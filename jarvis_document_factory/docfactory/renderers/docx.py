"""DOCX_Renderer: SPEC to DOCX via python-docx (generalizes build_docx.py).

Applies the PDF-style layout whether or not a PDF was also requested: A4, the
SPEC font at its size, the SPEC line spacing, justified body, one
page_break_before per chapter, and a PAGE field in the footer. When the SPEC
has a table of contents and a matching pdf_path is supplied, the TOC page
numbers are scanned from that PDF so the DOCX numbers match the PDF exactly.
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..images import verified_figure_refs
from ..spec import RenderResult, resolve_style
from .base import pinned_datetime, resolve_logo_path, slug

BLACK = RGBColor(0x00, 0x00, 0x00)
RULE_HEX = "000000"
TBL_HEAD_HEX = "ECECEC"
TBL_BORDER_HEX = "CCCCCC"
TBL_ZEBRA_HEX = "F7F7F7"


# ------------------------------------------------------------------ helpers
def _set_font(run, font, size=12, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), font)


def _ls(pf, spacing):
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spacing


def _flag(p, tag):
    pPr = p._p.get_or_add_pPr()
    e = OxmlElement("w:" + tag)
    e.set(qn("w:val"), "true")
    pPr.append(e)


def _bottom_rule(p, color=RULE_HEX, sz="6"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _page_field(paragraph, font):
    r = paragraph.add_run()
    e = OxmlElement("w:fldChar")
    e.set(qn("w:fldCharType"), "begin")
    r._r.append(e)
    _set_font(r, font, 11)
    r = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " PAGE "
    r._r.append(it)
    _set_font(r, font, 11)
    r = paragraph.add_run()
    e = OxmlElement("w:fldChar")
    e.set(qn("w:fldCharType"), "separate")
    r._r.append(e)
    _set_font(r, font, 11)
    r = paragraph.add_run("1")
    _set_font(r, font, 11)
    r = paragraph.add_run()
    e = OxmlElement("w:fldChar")
    e.set(qn("w:fldCharType"), "end")
    r._r.append(e)
    _set_font(r, font, 11)


def _enable_update_fields(doc):
    s = doc.settings.element
    if s.find(qn("w:updateFields")) is None:
        u = OxmlElement("w:updateFields")
        u.set(qn("w:val"), "true")
        s.append(u)


def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), fill)
    tcPr.append(sh)


def _set_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), TBL_BORDER_HEX)
        borders.append(e)
    tblPr.append(borders)


class _Builder:
    def __init__(self, doc, style):
        self.doc = doc
        self.font = style["font"]
        self.size = style["font_size_pt"]
        self.spacing = style["line_spacing"]
        self.align = {
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(style["body_align"], WD_ALIGN_PARAGRAPH.JUSTIFY)
        self.text_w_cm = 21 - style["margins_cm"]["left"] - style["margins_cm"]["right"]

    def para(self, text, size=None, after=6, before=0, align=None, bold=False):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        p.alignment = self.align if align is None else align
        _ls(pf, self.spacing)
        pf.space_after = Pt(after)
        pf.space_before = Pt(before)
        _flag(p, "widowControl")
        r = p.add_run(text)
        _set_font(r, self.font, size or self.size, bold=bold)
        return p

    def heading(self, text, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=16,
                page_break=False, italic=False, rule=False):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        p.alignment = align
        _ls(pf, self.spacing)
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.keep_with_next = True
        _flag(p, "widowControl")
        _flag(p, "keepLines")
        if page_break:
            pf.page_break_before = True
        r = p.add_run(text)
        _set_font(r, self.font, size, bold=True, italic=italic, color=BLACK)
        if rule:
            _bottom_rule(p)
        return p

    def callout(self, text):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _ls(pf, self.spacing)
        pf.space_before = Pt(4)
        pf.space_after = Pt(10)
        pf.left_indent = Cm(0.4)
        pf.right_indent = Cm(0.2)
        _flag(p, "widowControl")
        _flag(p, "keepLines")
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), TBL_ZEBRA_HEX)
        pPr.append(shd)
        pbdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), TBL_BORDER_HEX)
        pbdr.append(left)
        pPr.append(pbdr)
        r = p.add_run(text)
        _set_font(r, self.font, self.size)

    def listing(self, items, ordered):
        for i, it in enumerate(items, 1):
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _ls(pf, self.spacing)
            pf.space_after = Pt(4)
            pf.left_indent = Cm(1.0)
            pf.first_line_indent = Cm(-0.6)
            _flag(p, "widowControl")
            marker = ("%d. " % i) if ordered else "- "
            r = p.add_run(marker + str(it))
            _set_font(r, self.font, self.size)

    def table(self, tb):
        header = tb.get("header", [])
        if not header:
            return
        if tb.get("caption"):
            cap = self.doc.add_paragraph()
            cap.paragraph_format.space_after = Pt(3)
            cap.paragraph_format.space_before = Pt(6)
            cap.paragraph_format.keep_with_next = True
            _flag(cap, "widowControl")
            rc = cap.add_run(tb["caption"])
            _set_font(rc, self.font, 11, italic=True, color=BLACK)
        t = self.doc.add_table(rows=1, cols=len(header))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_borders(t)
        hdr = t.rows[0].cells
        for j, h in enumerate(header):
            _shade(hdr[j], TBL_HEAD_HEX)
            pr = hdr[j].paragraphs[0]
            pr.paragraph_format.space_after = Pt(2)
            rr = pr.add_run(str(h))
            _set_font(rr, self.font, 11, bold=True, color=BLACK)
        trPr = t.rows[0]._tr.get_or_add_trPr()
        th = OxmlElement("w:tblHeader")
        th.set(qn("w:val"), "true")
        trPr.append(th)
        for ri, row in enumerate(tb.get("rows", [])):
            cells = t.add_row().cells
            for j, c in enumerate(row):
                if j >= len(header):
                    break
                if ri % 2 == 1:
                    _shade(cells[j], TBL_ZEBRA_HEX)
                pr = cells[j].paragraphs[0]
                pr.paragraph_format.space_after = Pt(2)
                rr = pr.add_run(str(c))
                _set_font(rr, self.font, 11)
        for row in t.rows:
            rpr = row._tr.get_or_add_trPr()
            cs = OxmlElement("w:cantSplit")
            cs.set(qn("w:val"), "true")
            rpr.append(cs)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def figure(self, fig):
        pic = self.doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.paragraph_format.space_before = Pt(8)
        pic.paragraph_format.space_after = Pt(2)
        _flag(pic, "keepLines")
        pic.add_run().add_picture(fig.path, width=Cm(10))
        if fig.caption:
            cp = self.doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(12)
            rr = cp.add_run(fig.caption)
            _set_font(rr, self.font, 11, italic=True)

    def toc_line(self, label, page, lvl1=True):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.12
        pf.space_after = Pt(2 if lvl1 else 1)
        if not lvl1:
            pf.left_indent = Cm(0.9)
        pf.tab_stops.add_tab_stop(Cm(self.text_w_cm), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(label + "\t" + str(page))
        _set_font(r, self.font, self.size, bold=lvl1)


def _scan_pages(pdf_path, spec):
    """Map section/heading keys to printed page numbers in the reference PDF."""
    from pypdf import PdfReader

    pages = [(p.extract_text() or "") for p in PdfReader(pdf_path).pages]
    norm = [" ".join(t.split()) for t in pages]

    def find(label):
        key = " ".join(str(label).split())
        if not key:
            return ""
        for i, t in enumerate(norm):
            if key in t:
                return i + 1
        return ""

    pm = {}
    for s in spec.sections:
        if s.kind == "toc":
            continue
        if s.kind == "chapter" and s.number:
            pm[s.id] = find("BAB %s" % s.number)
        else:
            pm[s.id] = find(s.title)
        if s.kind == "chapter":
            for b in s.blocks:
                if b.get("type") == "heading" and b.get("level", 2) == 2:
                    hid = b.get("id") or slug(b.get("text", ""))
                    pm[hid] = find(b.get("text", ""))
    return pm


def _render_blocks(builder, spec, blocks, verified_refs):
    for b in blocks:
        t = b.get("type")
        if t == "heading":
            level = b.get("level", 2)
            builder.heading(b.get("text", ""), size=builder.size,
                            align=WD_ALIGN_PARAGRAPH.LEFT, before=10 if level == 2 else 8,
                            after=4 if level == 2 else 3, italic=(level == 3))
        elif t in ("paragraph", "lead"):
            builder.para(b.get("text", ""))
        elif t == "callout":
            builder.callout(b.get("text", ""))
        elif t == "list":
            builder.listing(b.get("items", []), b.get("ordered", False))
        elif t == "table":
            builder.table(b)
        elif t == "figure":
            ref = b.get("ref", "")
            fig = spec.figure_by_ref(ref)
            if fig is not None and ref in verified_refs:
                builder.figure(fig)


def _title_page(builder, spec):
    idn = spec.identity

    def tp(text, size, bold=False, italic=False, before=0, after=6, upper=False):
        p = builder.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        _ls(pf, builder.spacing)
        r = p.add_run(text.upper() if upper else text)
        _set_font(r, builder.font, size, bold=bold, italic=italic, color=BLACK)

    tp(idn.title, 18, bold=True, before=18, after=12, upper=True)
    if idn.subtitle:
        tp(idn.subtitle, 12.5, italic=True, before=4, after=20)
    if idn.course:
        tp("Disusun untuk memenuhi tugas mata kuliah", 12, after=2)
        tp(idn.course, 12, bold=True, after=8)
    if idn.lecturer:
        tp("Dosen Pengampu: " + idn.lecturer, 12, after=18)
    logo = resolve_logo_path(spec)
    if logo:
        lp = builder.doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(2)
        lp.paragraph_format.space_after = Pt(18)
        _flag(lp, "keepLines")
        lp.add_run().add_picture(logo, width=Cm(4.2))
    if idn.authors:
        tp("Disusun oleh:", 12, bold=True, after=8)
        for a in idn.authors:
            tp(a.get("name", ""), 12.5, bold=True, after=1)
            if a.get("id"):
                tp("NIM. " + a["id"], 11.5, after=8)
    for i, line in enumerate([x for x in (idn.program, idn.faculty, idn.institution, idn.year) if x]):
        tp(line, 12.5, bold=True, before=18 if i == 0 else 0, after=2, upper=True)


def render(spec, out_path: str, *, pdf_path: str | None = None) -> RenderResult:
    style = resolve_style(spec)
    verified_refs = verified_figure_refs(spec)
    m = style["margins_cm"]

    doc = Document()
    doc.styles["Normal"].font.name = style["font"]
    doc.styles["Normal"].font.size = Pt(style["font_size_pt"])
    builder = _Builder(doc, style)

    # Section 1: title page (no footer page number)
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(m["top"])
    sec.bottom_margin = Cm(m["bottom"])
    sec.left_margin = Cm(m["left"])
    sec.right_margin = Cm(m["right"])
    sec.footer.is_linked_to_previous = False
    _title_page(builder, spec)

    # Section 2: body with footer PAGE field
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width = Cm(21)
    body.page_height = Cm(29.7)
    body.top_margin = Cm(m["top"])
    body.bottom_margin = Cm(m["bottom"])
    body.left_margin = Cm(m["left"])
    body.right_margin = Cm(m["right"])
    body.footer.is_linked_to_previous = False
    fp = body.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _page_field(fp, style["font"])

    pm = _scan_pages(pdf_path, spec) if pdf_path else {}

    first_content = True
    for s in spec.sections:
        if s.kind == "toc":
            builder.heading(s.title, size=14, page_break=not first_content, rule=True)
            for t in spec.sections:
                if t.kind == "toc":
                    continue
                label = ("BAB %s  %s" % (t.number, t.title)) if (t.kind == "chapter" and t.number) else t.title
                builder.toc_line(label, pm.get(t.id, ""), lvl1=True)
                if t.kind == "chapter":
                    for b in t.blocks:
                        if b.get("type") == "heading" and b.get("level", 2) == 2:
                            hid = b.get("id") or slug(b.get("text", ""))
                            builder.toc_line(b.get("text", ""), pm.get(hid, ""), lvl1=False)
        elif s.kind == "chapter":
            if s.number:
                builder.heading("BAB %s" % s.number, size=14, page_break=True, after=4)
                builder.heading(s.title, size=14, after=16, rule=True)
            else:
                builder.heading(s.title, size=14, page_break=True, after=16, rule=True)
            _render_blocks(builder, spec, s.blocks, verified_refs)
        elif s.kind == "references":
            builder.heading(s.title, size=14, page_break=not first_content, after=16, rule=True)
            for r in spec.references:
                p = doc.add_paragraph()
                pf = p.paragraph_format
                _ls(pf, builder.spacing)
                pf.space_after = Pt(8)
                pf.left_indent = Cm(1.2)
                pf.first_line_indent = Cm(-1.2)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _flag(p, "widowControl")
                rr = p.add_run(r.apa)
                _set_font(rr, builder.font, builder.size)
        else:  # frontmatter, appendix
            builder.heading(s.title, size=14, page_break=not first_content, after=16, rule=True)
            _render_blocks(builder, spec, s.blocks, verified_refs)
        first_content = False

    _enable_update_fields(doc)

    # Pin core properties for reproducibility (no wall-clock timestamps).
    dt = pinned_datetime(spec.doc_id).replace(tzinfo=None)
    cp = doc.core_properties
    cp.created = dt
    cp.modified = dt
    cp.title = spec.identity.title
    cp.revision = 1

    out_dir = os.path.dirname(os.path.abspath(out_path))
    os.makedirs(out_dir, exist_ok=True)
    doc.save(out_path)

    # Page count is not known without a layout engine; report chapter count as a
    # structural proxy when no PDF reference is available.
    page_count = 0
    if pdf_path and os.path.exists(pdf_path):
        try:
            from pypdf import PdfReader

            page_count = len(PdfReader(pdf_path).pages)
        except Exception:
            page_count = 0
    return RenderResult(fmt="docx", out_path=out_path, page_count=page_count, warnings=[])
