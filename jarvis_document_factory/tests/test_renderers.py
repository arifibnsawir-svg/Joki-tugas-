"""Property + integration tests for the renderers and multi-format dispatch."""
from __future__ import annotations

import copy
import os
import re
import tempfile

from hypothesis import given, settings
from hypothesis import strategies as st

from docfactory.orchestrator import diff_consistency, extract_presence, render_all
from docfactory.renderers import docx as docx_renderer
from docfactory.renderers import pdf as pdf_renderer
from docfactory.renderers.pptx import resolve_pptx_style
from docfactory.readers import count_images, read_text
from docfactory.spec import parse_spec
from strategies import clean_text, phrase


def _prepare(data):
    """Apply layers + resolve figures + validate, returning the SPEC."""
    from docfactory.citation import apply_citation_layer
    from docfactory.humanizer import humanize_spec
    from docfactory.images import assert_referenced_images_exist, resolve_figures
    from docfactory.validation import validate

    spec = parse_spec(data)
    humanize_spec(spec)
    apply_citation_layer(spec)
    resolve_figures(spec)
    assert_referenced_images_exist(spec)
    validate(spec)
    return spec


@st.composite
def small_spec(draw, formats=None, n_chapters=None, with_toc=False):
    fmts = formats if formats is not None else draw(
        st.lists(st.sampled_from(["pdf", "docx", "pptx"]), min_size=1, max_size=3, unique=True))
    nch = n_chapters if n_chapters is not None else draw(st.integers(min_value=1, max_value=3))
    sections = []
    if with_toc:
        sections.append({"id": "toc", "kind": "toc", "title": "DAFTAR ISI", "blocks": []})
    roman = ["I", "II", "III", "IV", "V"]
    for i in range(nch):
        sections.append({
            "id": "bab%d" % i, "kind": "chapter", "number": roman[i],
            "title": draw(phrase),
            "blocks": [
                {"type": "heading", "level": 2, "id": "h%d" % i, "text": draw(phrase)},
                {"type": "paragraph", "text": draw(clean_text)},
            ],
        })
    data = {
        "doc_id": draw(clean_text), "formats": fmts, "is_academic": False,
        "identity": {"title": draw(phrase), "authors": [{"name": draw(phrase), "id": "1"}]},
        "style": {"page_size": "A4", "font": "Times New Roman"},
        "sections": sections, "references": [], "figures": [],
    }
    return data


# Feature: jarvis-document-factory, Property 1: One file per requested format
@settings(max_examples=100, deadline=None)
@given(small_spec())
def test_property_1_one_file_per_format(data):
    spec = _prepare(data)
    with tempfile.TemporaryDirectory() as d:
        res = render_all(spec, d, basename="doc")
        assert set(res["results"].keys()) == set(spec.formats)
        for fmt, r in res["results"].items():
            assert r.out_path.endswith("." + fmt)
            assert os.path.exists(r.out_path) and os.path.getsize(r.out_path) > 0


# Feature: jarvis-document-factory, Property 3: Cross-format consistency
@settings(max_examples=60, deadline=None)
@given(small_spec(formats=["pdf", "docx", "pptx"]))
def test_property_3_cross_format_consistency(data):
    spec = _prepare(data)
    with tempfile.TemporaryDirectory() as d:
        res = render_all(spec, d, basename="doc")
        assert res["consistency"]["unmatched"] == []
        presences = {f: extract_presence(f, res["results"][f].out_path, spec)
                     for f in res["results"]}
        titles = {p["title"] for p in presences.values()}
        orders = {tuple(p["section_titles"]) for p in presences.values()}
        assert len(titles) == 1 and len(orders) == 1


# Feature: jarvis-document-factory, Property 6: Reproducible byte-equivalent output
@settings(max_examples=100, deadline=None)
@given(small_spec(formats=["pdf"]))
def test_property_6_reproducible_pdf(data):
    spec = _prepare(data)
    with tempfile.TemporaryDirectory() as d:
        a = os.path.join(d, "a.pdf")
        b = os.path.join(d, "b.pdf")
        pdf_renderer.render(spec, a)
        pdf_renderer.render(spec, b)
        with open(a, "rb") as fa, open(b, "rb") as fb:
            assert fa.read() == fb.read()


# Feature: jarvis-document-factory, Property 7: Output reflects only SPEC content
@settings(max_examples=60, deadline=None)
@given(small_spec(formats=["pdf"], n_chapters=1))
def test_property_7_output_reflects_only_spec(data):
    spec_a = _prepare(copy.deepcopy(data))
    data_b = copy.deepcopy(data)
    # change exactly one paragraph text field to a unique sentinel
    data_b["sections"][-1]["blocks"][1]["text"] = "Zenithsentinel unique marker phrase."
    spec_b = _prepare(data_b)
    with tempfile.TemporaryDirectory() as d:
        pa = os.path.join(d, "a.pdf")
        pb = os.path.join(d, "b.pdf")
        pdf_renderer.render(spec_a, pa)
        pdf_renderer.render(spec_b, pb)
        ta = " ".join(read_text("pdf", pa).split())
        tb = " ".join(read_text("pdf", pb).split())
        assert "Zenithsentinel unique marker phrase." in tb
        assert "Zenithsentinel unique marker phrase." not in ta


# Feature: jarvis-document-factory, Property 8: DOCX chapter breaks and footer page field
@settings(max_examples=100, deadline=None)
@given(st.integers(min_value=1, max_value=5))
def test_property_8_docx_chapter_breaks_and_footer(n):
    import docx as _docx
    from docx.oxml.ns import qn

    data = {
        "doc_id": "x", "formats": ["docx"], "is_academic": False,
        "identity": {"title": "Judul", "authors": []},
        "style": {"page_size": "A4", "font": "Times New Roman"},
        "sections": [
            {"id": "b%d" % i, "kind": "chapter", "number": ["I", "II", "III", "IV", "V"][i],
             "title": "Bab %d" % i, "blocks": [{"type": "paragraph", "text": "Isi bab."}]}
            for i in range(n)
        ],
        "references": [], "figures": [],
    }
    spec = _prepare(data)
    with tempfile.TemporaryDirectory() as d:
        out = os.path.join(d, "doc.docx")
        docx_renderer.render(spec, out)
        doc = _docx.Document(out)
        breaks = 0
        for p in doc.paragraphs:
            pPr = p._p.find(qn("w:pPr"))
            has_break = pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None
            if has_break and p.text.startswith("BAB "):
                breaks += 1
        assert breaks == n
        # footer PAGE field present in the body section
        found_page_field = False
        for section in doc.sections:
            xml = section.footer._element.xml
            if "PAGE" in xml and "fldChar" in xml:
                found_page_field = True
        assert found_page_field


# Feature: jarvis-document-factory, Property 9: PPTX style resolution
MINIMAL_REQUESTS = ["minimal", "buat minimal", "minimal style", "plain formatting",
                    "plain style", "keep it plain", "gaya minimal"]
HOUSE_REQUESTS = ["", "default slides", "plain default slides", "standard deck",
                  "slide biasa", "buatkan presentasi", "designed deck"]


@settings(max_examples=100)
@given(st.one_of(st.sampled_from(MINIMAL_REQUESTS), st.sampled_from(HOUSE_REQUESTS)))
def test_property_9_pptx_style_resolution(request):
    resolved = resolve_pptx_style(request)
    explicit_minimal = ("minimal" in request.lower()) or (
        "plain" in request.lower() and "default" not in request.lower())
    assert (resolved == "minimal") == explicit_minimal


# Feature: jarvis-document-factory, Property 27: Renderers embed only verified images
@settings(max_examples=60, deadline=None)
@given(st.integers(min_value=0, max_value=3))
def test_property_27_embed_only_verified_images(n_used):
    with tempfile.TemporaryDirectory() as d:
        # n_used verified figures referenced by blocks + 1 verified but UNUSED figure
        figures = []
        blocks = []
        for i in range(n_used):
            path = os.path.join(d, "u%d.png" % i)
            _write_png(path, shade=40 * i)
            figures.append({"ref": "u%d" % i, "path": path, "caption": "cap", "verified_path": False})
            blocks.append({"type": "figure", "ref": "u%d" % i})
        unused = os.path.join(d, "unused.png")
        _write_png(unused, shade=250)
        figures.append({"ref": "unused", "path": unused, "caption": "x", "verified_path": False})

        data = {
            "doc_id": "img", "formats": ["pdf"], "is_academic": False,
            "identity": {"title": "T", "authors": []},
            "style": {"page_size": "A4", "font": "Times New Roman"},
            "sections": [{"id": "lamp", "kind": "appendix", "title": "Lampiran", "blocks": blocks}],
            "references": [], "figures": figures,
        }
        spec = _prepare(data)
        html = pdf_renderer.build_html(spec)
        # exactly the used (referenced + verified) figures are embedded; the
        # verified-but-unused figure is never embedded (subset of verified refs)
        assert html.count("data:image") == n_used


# ---- integration tests ----
def test_integration_pdf_css_shape():
    data = {
        "doc_id": "css", "formats": ["pdf"], "is_academic": False,
        "identity": {"title": "T", "authors": []},
        "style": {"page_size": "A4", "font": "Times New Roman"},
        "sections": [
            {"id": "toc", "kind": "toc", "title": "DAFTAR ISI", "blocks": []},
            {"id": "b1", "kind": "chapter", "number": "I", "title": "Bab",
             "blocks": [{"type": "paragraph", "text": "Isi."}]},
        ],
        "references": [], "figures": [],
    }
    spec = _prepare(data)
    html = pdf_renderer.build_html(spec)
    assert "target-counter" in html
    assert "counter-reset" not in html
    assert "counter-set" not in html
    assert "hyphens:none" in html or "hyphens: none" in html


def test_integration_pdf_before_docx_toc_numbers():
    data = {
        "doc_id": "toc", "formats": ["pdf", "docx"], "is_academic": False,
        "identity": {"title": "Judul Panjang", "authors": []},
        "style": {"page_size": "A4", "font": "Times New Roman"},
        "sections": [
            {"id": "kp", "kind": "frontmatter", "title": "KATA PENGANTAR",
             "blocks": [{"type": "paragraph", "text": "Pengantar singkat."}]},
            {"id": "toc", "kind": "toc", "title": "DAFTAR ISI", "blocks": []},
            {"id": "b1", "kind": "chapter", "number": "I", "title": "PENDAHULUAN",
             "blocks": [{"type": "paragraph", "text": "Isi bab satu."}]},
            {"id": "b2", "kind": "chapter", "number": "II", "title": "LANDASAN TEORI",
             "blocks": [{"type": "paragraph", "text": "Isi bab dua."}]},
        ],
        "references": [], "figures": [],
    }
    spec = _prepare(data)
    with tempfile.TemporaryDirectory() as d:
        res = render_all(spec, d, basename="doc")
        pdf_path = res["results"]["pdf"].out_path
        docx_path = res["results"]["docx"].out_path
        expected = docx_renderer._scan_pages(pdf_path, spec)
        import docx as _docx
        doc = _docx.Document(docx_path)
        # find TOC lines "BAB I  ... <num>" and compare parsed number to scan
        toc_numbers = {}
        for p in doc.paragraphs:
            m = re.match(r"BAB (I{1,3}|IV|V)\b.*\t(\d+)$", p.text)
            if m:
                toc_numbers[m.group(1)] = int(m.group(2))
        assert toc_numbers, "no TOC lines parsed"
        assert toc_numbers["I"] == expected["b1"]
        assert toc_numbers["II"] == expected["b2"]


def test_partial_consistency_report_names_unmatched():
    presences = {
        "pdf": {"title": True, "section_titles": ["A", "B"], "references_present": True},
        "docx": {"title": False, "section_titles": ["A", "B"], "references_present": True},
    }
    report = diff_consistency(presences)
    assert "title" in report["unmatched"]
    assert "section_order" in report["matched"]


def _write_png(path, shade=200):
    from PIL import Image

    Image.new("RGB", (8, 8), (shade % 256, shade % 256, shade % 256)).save(path, "PNG")
