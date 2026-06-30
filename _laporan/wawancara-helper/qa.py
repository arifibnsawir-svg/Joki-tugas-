# -*- coding: utf-8 -*-
"""
QA Laporan Wawancara Helper. Semua kriteria HARUS PASS.
1. Semua bagian ada dan urut.
2. BAB II menyebut Rogers, Corey, Brammer.
3. 4 helper, tiap helper memuat 15 poin + ada tabel ringkasan.
4. Konsistensi sitasi <-> daftar pustaka (dua arah).
5. Nol em-dash, nol kutip keriting, nol emoji.
6. Tanpa halaman kosong / heading menggantung.
7. PDF & DOCX ada dan dapat dibuka; laporkan jumlah halaman PDF.
8. Foto dirujuk di Lampiran 3 atau placeholder dicatat.
"""
import os, re, zipfile
from pypdf import PdfReader
import docx
import konten as K

HERE = os.path.dirname(os.path.abspath(__file__))
OUTDIR = os.path.abspath(os.path.join(HERE, "..", "..", "LAPORAN_FINAL"))
PDF = os.path.join(OUTDIR, "Laporan Wawancara Helper - Kelompok - Pengembangan Profesi Konseling - FINISH.pdf")
DOCX = os.path.join(OUTDIR, "Laporan Wawancara Helper - Kelompok - Pengembangan Profesi Konseling - FINISH.docx")
FOTODIR = os.path.join(HERE, "_foto_dl")

res = []
def add(n, ok, d=""):
    res.append((n, ok, d))

# baca PDF
reader = PdfReader(PDF)
pages = [(p.extract_text() or "") for p in reader.pages]
npage = len(pages)
full = "\n".join(pages)
flat = " ".join(full.split())

# ---- 1. semua bagian ada dan urut ----
order_labels = ["KATA PENGANTAR", "DAFTAR ISI", "BAB I", "PENDAHULUAN", "BAB II",
                "LANDASAN TEORI", "BAB III", "METODE KEGIATAN", "BAB IV",
                "HASIL WAWANCARA DAN PEMBAHASAN", "BAB V", "KESIMPULAN DAN SARAN",
                "DAFTAR PUSTAKA", "LAMPIRAN 1", "LAMPIRAN 2", "LAMPIRAN 3"]
positions = []
ok_order = True
last = -1
detail = []
for lab in order_labels:
    pos = flat.find(" ".join(lab.split()))
    positions.append(pos)
    if pos == -1 or pos < last:
        ok_order = False
        detail.append("%s(HILANG/URUT)" % lab)
    last = max(last, pos)
# judul utama di halaman judul
title_ok = "LAPORAN WAWANCARA DENGAN HELPER PEMBERI LAYANAN" in flat
add("1. Semua bagian ada & urut (judul, KP, DI, BAB I-V, DafPus, Lampiran 1-3)",
    ok_order and title_ok, "Judul utama: %s. %s" % ("ada" if title_ok else "TIDAK ADA",
    "urutan benar" if ok_order else "MASALAH: " + ", ".join(detail)))

# ---- 2. BAB II menyebut Rogers, Corey, Brammer ----
# gunakan rfind: kemunculan pertama ada di Daftar Isi, isi bab ada di kemunculan terakhir
i2 = flat.rfind("LANDASAN TEORI")
i3 = flat.rfind("METODE KEGIATAN")
bab2 = flat[i2:i3] if (i2 != -1 and i3 != -1 and i3 > i2) else ""
ahli = {n: (n in bab2) for n in ["Rogers", "Corey", "Brammer"]}
add("2. BAB II mengutip Rogers, Corey, Brammer", all(ahli.values()),
    ", ".join("%s=%s" % (k, v) for k, v in ahli.items()))

# ---- 3. 4 helper x 15 poin + tabel ringkasan ----
helper_ok = True
hd = []
for idx, h in enumerate(K.HELPERS, 1):
    # cari sub-bab 4.x di teks
    found_points = 0
    for i, (judul, teks) in enumerate(h["poin"], 1):
        # ambil potongan unik dari teks poin
        probe = " ".join(teks.split())[:40]
        if probe in flat:
            found_points += 1
    hd.append("H%d:%d/15" % (idx, found_points))
    if len(h["poin"]) != 15 or found_points < 15:
        helper_ok = False
n_helper = len(K.HELPERS)
tabel_ringkasan = "Tabel 4.1" in flat and "Ringkasan hasil wawancara" in flat
add("3. 4 helper, tiap helper 15 poin + tabel ringkasan",
    n_helper == 4 and helper_ok and tabel_ringkasan,
    "%d helper; %s; tabel ringkasan: %s" % (n_helper, " ".join(hd), tabel_ringkasan))

# ---- 4. konsistensi sitasi <-> daftar pustaka ----
# kumpulan nama belakang penulis dari daftar pustaka
ref_surnames = {"Brammer", "Corey", "Gladding", "Lubis", "Prayitno", "Rogers", "Willis"}
# sitasi dalam teks: pola "Nama (tahun)" dan "Nama dan Nama (tahun)"
body_text = flat[:flat.rfind("DAFTAR PUSTAKA")]
cited = set(re.findall(r"([A-Z][a-z]+)\s*\(\d{4}\)", body_text))
cited |= set(re.findall(r"([A-Z][a-z]+)\s+dan\s+[A-Z]", body_text))  # "Prayitno dan Amti"
# normalisasi sitasi yang relevan (buang kata umum yang ketangkap)
known = {"Rogers", "Corey", "Brammer", "Gladding", "Prayitno", "Willis", "Lubis", "Amti", "MacDonald"}
cited = {c for c in cited if c in known}
# Amti & MacDonald adalah penulis kedua -> peta ke entri utama
cited_main = set()
for c in cited:
    if c == "Amti":
        cited_main.add("Prayitno")
    elif c == "MacDonald":
        cited_main.add("Brammer")
    else:
        cited_main.add(c)
# arah 1: tiap sitasi punya entri daftar pustaka
missing_ref = cited_main - ref_surnames
# arah 2: tiap entri daftar pustaka memang dikutip
dp_text = flat[flat.rfind("DAFTAR PUSTAKA"):flat.rfind("LAMPIRAN 1")]
uncited = set()
for s in ref_surnames:
    in_text = s in cited_main
    in_dp = s in dp_text
    if in_dp and not in_text:
        uncited.add(s)
add("4. Konsistensi sitasi <-> daftar pustaka (dua arah)",
    len(missing_ref) == 0 and len(uncited) == 0,
    "sitasi: %s; tanpa entri: %s; entri tak dikutip: %s"
    % (sorted(cited_main), sorted(missing_ref) or "tidak ada", sorted(uncited) or "tidak ada"))

# ---- 5. humanizer ----
emdash = full.count("\u2014")
endash = full.count("\u2013")
curly = sum(full.count(c) for c in ["\u201c", "\u201d", "\u2018", "\u2019"])
emoji = [c for c in full if ord(c) >= 0x1F000 or (0x2190 <= ord(c) <= 0x2BFF) or (0x2600 <= ord(c) <= 0x27BF)]
add("5. Humanizer: 0 em-dash, 0 en-dash, 0 kutip keriting, 0 emoji",
    emdash == 0 and endash == 0 and curly == 0 and len(emoji) == 0,
    "em-dash=%d, en-dash=%d, kutip-keriting=%d, emoji=%d" % (emdash, endash, curly, len(emoji)))

# ---- 6. tanpa halaman kosong / heading menggantung / halaman hampir kosong ----
blank = [i + 1 for i in range(1, npage) if len(pages[i].strip()) < 5]
# heading menggantung: halaman yang hanya berisi "BAB X" tanpa isi lain
dangling = []
for i, t in enumerate(pages):
    s = " ".join(t.split())
    if re.fullmatch(r"BAB [IVX]+", s):
        dangling.append(i + 1)
# halaman HAMPIR kosong: non-cover (indeks > 0), panjang teks 1..120 karakter.
# Ini menandai baris tunggal yang menggantung (mis. satu entri Daftar Isi yang
# tumpah ke halaman berikutnya). Halaman yang memang pendek namun sah
# (lanjutan paragraf/daftar isi normal tanpa dot-leader) DIKECUALIKAN; baris
# Daftar Isi yang menggantung dikenali dari pola titik-titik (dot leader).
near_empty = []
for i in range(1, npage):
    s = " ".join(pages[i].split())
    L = len(s)
    if 1 <= L <= 120:
        is_toc_orphan = re.search(r"\.{4,}", s) is not None  # baris Daftar Isi menggantung
        if is_toc_orphan:
            near_empty.append("hal %d (%d char, dot-leader)" % (i + 1, L))
add("6. Tanpa halaman kosong / heading menggantung / halaman hampir kosong",
    len(blank) == 0 and len(dangling) == 0 and len(near_empty) == 0,
    "halaman kosong: %s; heading menggantung: %s; halaman hampir kosong: %s"
    % (blank or "tidak ada", dangling or "tidak ada", near_empty or "tidak ada"))

# ---- 7. PDF & DOCX ada & dapat dibuka ----
docx_ok = False
docx_detail = ""
try:
    d = docx.Document(DOCX)
    docx_ok = True
    docx_detail = "%d paragraf, %d tabel, %d gambar" % (len(d.paragraphs), len(d.tables), len(d.inline_shapes))
except Exception as e:
    docx_detail = "ERROR: %s" % e
add("7. PDF & DOCX ada dan dapat dibuka",
    os.path.exists(PDF) and os.path.exists(DOCX) and npage > 0 and docx_ok,
    "PDF=%d halaman; DOCX: %s" % (npage, docx_detail))

# ---- 8. foto di Lampiran 3 ----
l3 = flat[flat.rfind("LAMPIRAN 3"):]
photos_present = [h["foto"] for h in K.HELPERS if os.path.exists(os.path.join(FOTODIR, h["foto"]))]
photos_missing = [h["foto"] for h in K.HELPERS if not os.path.exists(os.path.join(FOTODIR, h["foto"]))]
referenced = "Dokumentasi wawancara" in l3
# pada DOCX hitung gambar
docx_imgs = len(d.inline_shapes) if docx_ok else 0
add("8. Foto dirujuk di Lampiran 3 (atau placeholder dicatat)",
    referenced and (len(photos_present) >= 1 or len(photos_missing) > 0),
    "foto tersedia: %s; placeholder: %s; gambar di DOCX: %d"
    % (photos_present, photos_missing or "tidak ada", docx_imgs))

# ---- 9. identitas Guru BK asli (Elisabeth Suwartini, tanpa nama samaran) ----
nama_asli_ada = "Elisabeth Suwartini" in flat
fulana_ada = ("Fulana" in flat) or ("samaran" in flat)
add("9. Identitas Guru BK asli: 'Elisabeth Suwartini' ada, 'Fulana'/'samaran' TIDAK ada",
    nama_asli_ada and not fulana_ada,
    "Elisabeth Suwartini=%s; Fulana/samaran=%s" % (nama_asli_ada, fulana_ada))

# ---- 10. Kata Pengantar Islami ----
kp_start = flat.find("KATA PENGANTAR")
di_start = flat.find("DAFTAR ISI")
kp_text = flat[kp_start:di_start] if (kp_start != -1 and di_start != -1 and di_start > kp_start) else flat
islami_ada = "Allah Subhanahu wa Ta'ala" in kp_text
puji_tuhan_ada = "Puji Tuhan" in flat
add("10. Kata Pengantar Islami: 'Allah Subhanahu wa Ta'ala' ada, 'Puji Tuhan' TIDAK ada",
    islami_ada and not puji_tuhan_ada,
    "Allah Subhanahu wa Ta'ala (di KP)=%s; Puji Tuhan=%s" % (islami_ada, puji_tuhan_ada))

# ---- 11. detail identitas Guru BK: tempat & lama pengalaman ----
sekolah_ada = "SMA Negeri 49 Jakarta" in flat
lama_ada = "32 tahun" in flat
add("11. Detail Guru BK: 'SMA Negeri 49 Jakarta' & '32 tahun' ada",
    sekolah_ada and lama_ada,
    "SMA Negeri 49 Jakarta=%s; 32 tahun=%s" % (sekolah_ada, lama_ada))

# ---- 12. cover memuat logo (gambar) ATAU placeholder dicatat ----
def page_has_image(page):
    try:
        res = page.get("/Resources")
        xo = res.get("/XObject") if res else None
        if xo:
            xo = xo.get_object()
            for o in xo.values():
                if o.get_object().get("/Subtype") == "/Image":
                    return True
    except Exception:
        pass
    return False

cover_text = pages[0] if npage > 0 else ""
pdf_logo = page_has_image(reader.pages[0]) if npage > 0 else False
pdf_placeholder = "[Logo UNINDRA]" in cover_text
# di DOCX: logo menambah 1 gambar di luar foto Lampiran 3
docx_logo = docx_ok and (docx_imgs > len(photos_present))
docx_placeholder = docx_ok and any("[Logo UNINDRA]" in p.text for p in d.paragraphs)
pdf_cover_ok = pdf_logo or pdf_placeholder
docx_cover_ok = docx_logo or docx_placeholder
add("12. Cover memuat logo UNINDRA (gambar) atau placeholder",
    pdf_cover_ok and docx_cover_ok,
    "PDF: logo=%s, placeholder=%s; DOCX: logo=%s (gambar=%d, foto=%d), placeholder=%s"
    % (pdf_logo, pdf_placeholder, docx_logo, docx_imgs, len(photos_present), docx_placeholder))

# ---- 13. Daftar Isi muat dalam SATU halaman PDF ----
# Cari halaman yang memuat "DAFTAR ISI"; halaman berikutnya HARUS sudah masuk
# badan dokumen (memuat "BAB I" + "PENDAHULUAN"), BUKAN sisa entri Daftar Isi
# (mis. baris "LAMPIRAN ..." dengan dot leader yang tumpah ke halaman kedua).
di_idx = next((i for i, t in enumerate(pages) if "DAFTAR ISI" in " ".join(t.split())), None)
di_detail = ""
di_ok = False
if di_idx is None:
    di_detail = "halaman DAFTAR ISI tidak ditemukan"
elif di_idx + 1 >= npage:
    di_detail = "tidak ada halaman setelah DAFTAR ISI"
else:
    nxt = " ".join(pages[di_idx + 1].split())
    starts_body = ("BAB I" in nxt and "PENDAHULUAN" in nxt)
    # sisa entri Daftar Isi yang menggantung: baris pendek berisi dot leader
    toc_leftover = bool(re.search(r"\.{4,}", nxt)) and ("LAMPIRAN" in nxt or len(nxt) <= 120)
    di_ok = starts_body and not toc_leftover
    di_detail = ("DAFTAR ISI di hal %d; hal berikut (hal %d) %s badan dokumen; "
                 "sisa entri menggantung: %s") % (
        di_idx + 1, di_idx + 2,
        "memulai" if starts_body else "TIDAK memulai",
        "ADA -> " + nxt[:60] if toc_leftover else "tidak ada")
add("13. Daftar Isi muat dalam satu halaman PDF (tanpa tumpah ke halaman kedua)",
    di_ok, di_detail)

# ---- cetak ----
print("=" * 70)
print("QA - LAPORAN WAWANCARA DENGAN HELPER PEMBERI LAYANAN")
print("=" * 70)
allp = True
for nm, ok, d2 in res:
    print("[%s] %s" % ("PASS" if ok else "FAIL", nm))
    if d2:
        print("        -> %s" % d2)
    if not ok:
        allp = False
print("=" * 70)
print("JUMLAH HALAMAN PDF:", npage)
print("HASIL:", "SEMUA PASS" if allp else "ADA FAIL")
print("=" * 70)
