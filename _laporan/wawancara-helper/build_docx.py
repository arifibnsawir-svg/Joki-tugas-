# -*- coding: utf-8 -*-
"""
Bangun versi DOCX Laporan Wawancara Helper (mirror layout PDF).
A4, Times New Roman 12, spasi 1.5, justify, margin kiri 4 / kanan 3 / atas 3 / bawah 3 cm.
page_break_before per BAB (tanpa halaman kosong), footer berisi field PAGE,
Daftar Isi MANUAL dengan nomor halaman hasil scan PDF acuan.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import konten as K

HERE = os.path.dirname(os.path.abspath(__file__))
OUTDIR = os.path.abspath(os.path.join(HERE, "..", "..", "LAPORAN_FINAL"))
OUTDOCX = os.path.join(OUTDIR, "Laporan Wawancara Helper - Kelompok - Pengembangan Profesi Konseling - FINISH.docx")
OUTPDF = os.path.join(OUTDIR, "Laporan Wawancara Helper - Kelompok - Pengembangan Profesi Konseling - FINISH.pdf")
FOTODIR = os.path.join(HERE, "_foto_dl")
ASSETDIR = os.path.join(HERE, "assets")
LOGO = os.path.join(ASSETDIR, "logo-unindra.jpg")
FONT = "Times New Roman"
TEXT_W_CM = 14.0  # lebar area teks A4 (21 - 4 - 3)
BLACK = RGBColor(0x00, 0x00, 0x00)      # seluruh teks hitam
NAVY = BLACK                            # alias: tidak ada lagi warna navy pada teks
NAVY_DK = BLACK
RULE_HEX = "000000"                     # garis bawah judul: hitam tipis
TBL_HEAD_HEX = "ECECEC"                 # header tabel abu-abu lembut
TBL_BORDER_HEX = "CCCCCC"               # garis tabel abu-abu tipis
TBL_ZEBRA_HEX = "F7F7F7"                # zebra sangat samar

# ---------------------------------------------------------------- scan PDF
def scan_pages():
    """Petakan label Daftar Isi -> nomor halaman tercetak di PDF acuan."""
    from pypdf import PdfReader
    pages = [(p.extract_text() or "") for p in PdfReader(OUTPDF).pages]
    norm = [" ".join(t.split()) for t in pages]

    def find(label):
        key = " ".join(label.split())
        for i, t in enumerate(norm):
            if key in t:
                return i + 1  # nomor tercetak = indeks 1-based
        return ""

    pm = {"KATA PENGANTAR": find("KATA PENGANTAR")}
    for b in K.BAB_LIST:
        pm["bab" + b["no"]] = find("BAB %s" % b["no"])
        for blk in b["isi"]:
            if blk[0] == "h2":
                pm[blk[1]] = find(blk[1])
    pm["DAFTAR PUSTAKA"] = find("DAFTAR PUSTAKA")
    pm["lamp1"] = find("LAMPIRAN 1")
    pm["lamp2"] = find("LAMPIRAN 2")
    pm["lamp3"] = find("LAMPIRAN 3")
    return pm

# ---------------------------------------------------------------- helpers
def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), FONT)

def ls15(pf):
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5

def ppr_flag(p, tag):
    pPr = p._p.get_or_add_pPr()
    e = OxmlElement("w:" + tag)
    e.set(qn("w:val"), "true")
    pPr.append(e)

def para(doc, text, size=12, after=6, before=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    p.alignment = align
    ls15(pf)
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    ppr_flag(p, "widowControl")
    r = p.add_run(text)
    set_font(r, size, bold=bold)
    return p

def para_bottom_rule(p, color=RULE_HEX, sz="6"):
    """Garis bawah tipis pada paragraf (judul BAB / bagian) - hitam, halus."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)

def heading(doc, text, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=16, page_break=False, italic=False, color=NAVY, rule=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    p.alignment = align
    ls15(pf)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.keep_with_next = True
    ppr_flag(p, "widowControl")
    ppr_flag(p, "keepLines")
    if page_break:
        pf.page_break_before = True
    r = p.add_run(text)
    set_font(r, size, bold=True, italic=italic, color=color)
    if rule:
        para_bottom_rule(p)
    return p

def callout(doc, text):
    """Kotak kutipan motivasi: shading tipis + garis aksen di kiri."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ls15(pf)
    pf.space_before = Pt(4); pf.space_after = Pt(10)
    pf.left_indent = Cm(0.4); pf.right_indent = Cm(0.2)
    ppr_flag(p, "widowControl"); ppr_flag(p, "keepLines")
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), TBL_ZEBRA_HEX); pPr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left"); left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18"); left.set(qn("w:space"), "6"); left.set(qn("w:color"), TBL_BORDER_HEX)
    pbdr.append(left); pPr.append(pbdr)
    r = p.add_run(text); set_font(r, 12)
    return p

def numbered(doc, items):
    for i, it in enumerate(items, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ls15(pf)
        pf.space_after = Pt(4)
        pf.left_indent = Cm(1.0)
        pf.first_line_indent = Cm(-0.6)
        ppr_flag(p, "widowControl")
        r = p.add_run("%d. %s" % (i, it))
        set_font(r, 12)

def page_field(paragraph):
    r = paragraph.add_run(); e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "begin"); r._r.append(e); set_font(r, 11)
    r = paragraph.add_run(); it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "; r._r.append(it); set_font(r, 11)
    r = paragraph.add_run(); e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "separate"); r._r.append(e); set_font(r, 11)
    r = paragraph.add_run("2"); set_font(r, 11)
    r = paragraph.add_run(); e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end"); r._r.append(e); set_font(r, 11)

def enable_update_fields(doc):
    s = doc.settings.element
    if s.find(qn("w:updateFields")) is None:
        u = OxmlElement("w:updateFields"); u.set(qn("w:val"), "true"); s.append(u)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), fill); tcPr.append(sh)

def set_borders(table):
    tblPr = table._tbl.tblPr; borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge); e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4"); e.set(qn("w:color"), TBL_BORDER_HEX); borders.append(e)
    tblPr.append(borders)

def add_table(doc, tb):
    cap = doc.add_paragraph(); cap.paragraph_format.space_after = Pt(3); cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.keep_with_next = True
    ppr_flag(cap, "widowControl")
    rc = cap.add_run(tb["judul"]); set_font(rc, 11, italic=True, color=BLACK)
    t = doc.add_table(rows=1, cols=len(tb["head"])); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t)
    hdr = t.rows[0].cells
    for j, h in enumerate(tb["head"]):
        shade(hdr[j], TBL_HEAD_HEX); pr = hdr[j].paragraphs[0]; pr.paragraph_format.space_after = Pt(2)
        rr = pr.add_run(h); set_font(rr, 11, bold=True, color=BLACK)
    trPr = t.rows[0]._tr.get_or_add_trPr(); th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); trPr.append(th)
    for ri, row in enumerate(tb["rows"]):
        cells = t.add_row().cells
        for j, c in enumerate(row):
            if ri % 2 == 1:
                shade(cells[j], TBL_ZEBRA_HEX)
            pr = cells[j].paragraphs[0]; pr.paragraph_format.space_after = Pt(2)
            rr = pr.add_run(c); set_font(rr, 11)
    for row in t.rows:
        rpr = row._tr.get_or_add_trPr(); cs = OxmlElement("w:cantSplit"); cs.set(qn("w:val"), "true"); rpr.append(cs)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def toc_line(doc, label, page, lvl1=True):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.3
    pf.space_after = Pt(5 if lvl1 else 2)
    if not lvl1:
        pf.left_indent = Cm(0.9)
    pf.tab_stops.add_tab_stop(Cm(TEXT_W_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(label + "\t" + str(page)); set_font(r, 12, bold=lvl1)

def render_blocks(doc, blocks):
    for b in blocks:
        kind = b[0]
        if kind == "h2":
            heading(doc, b[1], size=12, align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=4, color=NAVY)
        elif kind == "h3":
            heading(doc, b[1], size=12, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=3, italic=True, color=NAVY_DK)
        elif kind == "p":
            para(doc, b[1])
        elif kind == "lead":
            para(doc, b[1])
        elif kind == "callout":
            callout(doc, b[1])
        elif kind == "num":
            numbered(doc, b[1])
        elif kind == "table":
            add_table(doc, b[1])

# ---------------------------------------------------------------- build
def build():
    pm = scan_pages()
    doc = Document()
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(12)

    # ---- Section 1: HALAMAN JUDUL (tanpa nomor halaman) ----
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3); sec.bottom_margin = Cm(3); sec.left_margin = Cm(4); sec.right_margin = Cm(3)
    sec.footer.is_linked_to_previous = False

    I = K.IDENTITAS
    def tp(text, size, bold=False, italic=False, before=0, after=6, upper=False, color=None, rule=False):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format; pf.space_before = Pt(before); pf.space_after = Pt(after); ls15(pf)
        r = p.add_run(text.upper() if upper else text); set_font(r, size, bold=bold, italic=italic, color=color or BLACK)
        if rule:
            para_bottom_rule(p)
        return p
    # Cover sederhana: judul, subjudul, mata kuliah, dosen, LOGO (tengah),
    # "Disusun oleh" + nama/NIM, lalu prodi/fakultas/universitas/tahun. Semua teks hitam.
    tp(I["judul"], 18, bold=True, before=18, after=12, upper=True)
    tp(I["subjudul"], 12.5, italic=True, before=4, after=20)
    tp("Disusun untuk memenuhi tugas mata kuliah", 12, after=2)
    tp(I["matakuliah"], 12, bold=True, after=8)
    tp("Dosen Pengampu: " + I["dosen"], 12, after=18)
    # Logo UNINDRA di tengah (atau placeholder bila gagal diunduh)
    logo_missing = False
    if os.path.exists(LOGO):
        lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(2); lp.paragraph_format.space_after = Pt(18)
        ppr_flag(lp, "keepLines")
        lp.add_run().add_picture(LOGO, width=Cm(4.2))
    else:
        logo_missing = True
        tp("[Logo UNINDRA]", 12, before=18, after=18)
    tp("Disusun oleh:", 12, bold=True, after=8)
    for n, nim in I["penyusun"]:
        tp(n, 12.5, bold=True, after=1)
        tp("NIM. " + nim, 11.5, after=8)
    tp(I["prodi"], 12.5, bold=True, before=18, after=2, upper=True)
    tp(I["fakultas"], 12.5, bold=True, after=2, upper=True)
    tp(I["universitas"], 12.5, bold=True, after=2, upper=True)
    tp(I["tahun"], 12.5, bold=True, after=2)

    # ---- Section 2: ISI (footer PAGE field) ----
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width = Cm(21); body.page_height = Cm(29.7)
    body.top_margin = Cm(3); body.bottom_margin = Cm(3); body.left_margin = Cm(4); body.right_margin = Cm(3)
    body.footer.is_linked_to_previous = False
    fp = body.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_field(fp)

    # KATA PENGANTAR
    heading(doc, "KATA PENGANTAR", rule=True)
    for t in K.KATA_PENGANTAR:
        para(doc, t)

    # DAFTAR ISI
    heading(doc, "DAFTAR ISI", page_break=True, rule=True)
    toc_line(doc, "KATA PENGANTAR", pm["KATA PENGANTAR"], lvl1=True)
    for b in K.BAB_LIST:
        toc_line(doc, "BAB %s  %s" % (b["no"], b["judul"]), pm["bab" + b["no"]], lvl1=True)
        for blk in b["isi"]:
            if blk[0] == "h2":
                toc_line(doc, blk[1], pm.get(blk[1], ""), lvl1=False)
    toc_line(doc, "DAFTAR PUSTAKA", pm["DAFTAR PUSTAKA"], lvl1=True)
    toc_line(doc, "LAMPIRAN 1: Daftar Pertanyaan Wawancara", pm["lamp1"], lvl1=True)
    toc_line(doc, "LAMPIRAN 2: Transkrip Wawancara", pm["lamp2"], lvl1=True)
    toc_line(doc, "LAMPIRAN 3: Dokumentasi Kegiatan", pm["lamp3"], lvl1=True)

    # BAB I - V
    for b in K.BAB_LIST:
        heading(doc, "BAB %s" % b["no"], size=14, page_break=True, after=4)
        heading(doc, b["judul"], size=14, after=16, rule=True)
        render_blocks(doc, b["isi"])

    # DAFTAR PUSTAKA
    heading(doc, "DAFTAR PUSTAKA", size=14, page_break=True, after=16, rule=True)
    for ref in K.DAFTAR_PUSTAKA:
        p = doc.add_paragraph(); pf = p.paragraph_format; ls15(pf)
        pf.space_after = Pt(8); pf.left_indent = Cm(1.2); pf.first_line_indent = Cm(-1.2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ppr_flag(p, "widowControl")
        r = p.add_run(ref); set_font(r, 12)

    # LAMPIRAN 1
    heading(doc, "LAMPIRAN 1: DAFTAR PERTANYAAN WAWANCARA", size=14, page_break=True, after=12, rule=True)
    para(doc, "Berikut lima belas butir pertanyaan yang disusun berdasarkan Rencana Pembelajaran Semester dan diajukan kepada seluruh narasumber dengan penyesuaian bahasa sesuai profesi.")
    numbered(doc, K.PERTANYAAN)

    # LAMPIRAN 2
    heading(doc, "LAMPIRAN 2: TRANSKRIP WAWANCARA", size=14, page_break=True, after=12, rule=True)
    for h in K.HELPERS:
        heading(doc, "Transkrip %s: %s (%s)" % (h["kode"], h["nama"], h["profesi"]),
                size=12, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=4)
        for i, (judul, teks) in enumerate(h["poin"], 1):
            para(doc, "%d. %s" % (i, K.PERTANYAAN[i - 1]), after=1, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
            para(doc, teks, after=6)

    # LAMPIRAN 3 (foto)
    heading(doc, "LAMPIRAN 3: DOKUMENTASI KEGIATAN", size=14, page_break=True, after=12, rule=True)
    para(doc, "Berikut dokumentasi kegiatan wawancara bersama keempat narasumber.")
    missing = []
    for h in K.HELPERS:
        path = os.path.join(FOTODIR, h["foto"])
        cap = "Dokumentasi wawancara dengan %s (%s)" % (h["nama"], h["profesi"])
        if os.path.exists(path):
            pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.paragraph_format.space_before = Pt(8); pic.paragraph_format.space_after = Pt(2)
            ppr_flag(pic, "keepLines")
            pic.add_run().add_picture(path, width=Cm(10))
        else:
            missing.append(h["foto"])
            para(doc, "[Foto belum tersedia]", align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(12)
        rr = cp.add_run(cap); set_font(rr, 11, italic=True)

    enable_update_fields(doc)
    os.makedirs(OUTDIR, exist_ok=True)
    doc.save(OUTDOCX)
    if logo_missing:
        print("PERINGATAN: logo UNINDRA tidak ditemukan ->", LOGO, "(placeholder dipakai)")
    if missing:
        print("PERINGATAN: foto tidak ditemukan ->", missing)
    print("SAVED:", OUTDOCX, os.path.getsize(OUTDOCX), "bytes")

if __name__ == "__main__":
    build()
