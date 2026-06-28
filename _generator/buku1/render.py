# -*- coding: utf-8 -*-
"""
Builder mini book PKN - Nurul Syifa (Buku #1).
Pertahankan cover depan (image1) + biografi belakang (image2) dari file asli.
Rebuild: Kata Pengantar (preserved) + Daftar Isi (TOC field) + Bab 1-8 + Daftar Pustaka.
DNA: naratif-reflektif. Gaya sudah di-humanize (variasi ritme, hindari pola AI).
"""
import os, zipfile, copy
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "drive_buku/Joki buku pkn/5_6183913965983112173.docx"
OUT = "Joki-tugas-/buku/Mini Book PKN - Nurul Syifa.docx"

from konten import KATA_PENGANTAR, BAB, DAFTAR_PUSTAKA, PENULIS

# ---------- ekstrak gambar cover asli ----------
def extract_covers():
    os.makedirs("build_nurul/img", exist_ok=True)
    z = zipfile.ZipFile(SRC)
    media = sorted([n for n in z.namelist() if n.startswith("word/media/image")])
    paths = []
    for i, m in enumerate(media, 1):
        ext = os.path.splitext(m)[1]
        p = f"build_nurul/img/cover{i}{ext}"
        open(p, "wb").write(z.read(m))
        paths.append(p)
    return paths  # [front, back]

# ---------- helper field (TOC, PAGE) ----------
def add_field(paragraph, instr):
    r = paragraph.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    fs = OxmlElement('w:fldChar'); fs.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = ""
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    r._r.append(fb); r._r.append(it); r._r.append(fs); r._r.append(t); r._r.append(fe)

def set_cell_or_page_size(section, w_cm, h_cm):
    section.page_width = Cm(w_cm); section.page_height = Cm(h_cm)

def full_page_image_section(doc, img, first=False):
    if first:
        sec = doc.sections[0]
    else:
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Cm(0)
    sec.header_distance = Cm(0); sec.footer_distance = Cm(0)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(img, width=Cm(21), height=Cm(29.7))
    return sec

def style_base(doc):
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'
    st.font.size = Pt(12)
    rpr = st.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    rf.set(qn('w:eastAsia'), 'Times New Roman')
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(10)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, sz, bold, sb, sa in [('Heading 1', 16, True, 18, 8), ('Heading 2', 13, True, 12, 6), ('Heading 3', 12, True, 8, 4)]:
        s = doc.styles[name]
        s.font.name = 'Times New Roman'; s.font.size = Pt(sz); s.font.bold = bold
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(sb); s.paragraph_format.space_after = Pt(sa)
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def add_page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, "PAGE")
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def heading(doc, text, level):
    h = doc.add_heading(level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT if level >= 2 else WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    return h

def para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.italic = italic
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    return p

def page_break(doc):
    doc.add_page_break()

# ---------- build ----------
def build():
    front, back = extract_covers()
    doc = Document()
    style_base(doc)

    # --- Section 1: COVER DEPAN (gambar asli) ---
    full_page_image_section(doc, front, first=True)

    # --- Section 2: ISI (margin wajar + nomor halaman) ---
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width = Cm(21); body.page_height = Cm(29.7)
    body.left_margin = Cm(3); body.right_margin = Cm(2.5)
    body.top_margin = Cm(3); body.bottom_margin = Cm(2.5)
    add_page_number_footer(body)

    # KATA PENGANTAR (dipertahankan)
    heading(doc, "KATA PENGANTAR", 1)
    for t in KATA_PENGANTAR:
        para(doc, t)
    page_break(doc)

    # DAFTAR ISI (TOC field - auto update di Word)
    heading(doc, "DAFTAR ISI", 1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-2" \\h \\z \\u')
    note = para(doc, "(Klik kanan daftar isi ini di Microsoft Word lalu pilih \"Update Field\" untuk memunculkan nomor halaman.)", italic=True)
    note.runs[0].font.size = Pt(10)
    page_break(doc)

    # BAB 1-8
    for bab in BAB:
        heading(doc, f"BAB {bab['no']}", 1)
        heading(doc, bab['judul'], 1)
        for blk in bab['isi']:
            kind, text = blk
            if kind == 'h2':
                heading(doc, text, 2)
            elif kind == 'h3':
                heading(doc, text, 3)
            else:
                para(doc, text)
        page_break(doc)

    # DAFTAR PUSTAKA
    heading(doc, "DAFTAR PUSTAKA", 1)
    for ref in DAFTAR_PUSTAKA:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.left_indent = Cm(1); pf.first_line_indent = Cm(-1)  # hanging indent
        r = p.add_run(ref)
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)

    # --- Section 3: COVER BELAKANG (biografi penulis - gambar asli) ---
    full_page_image_section(doc, back, first=False)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("SAVED:", OUT)

if __name__ == "__main__":
    build()
