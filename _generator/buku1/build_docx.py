# -*- coding: utf-8 -*-
"""DOCX rapi, layout meniru render PDF (A4, TNR 12, spasi 1.75) agar paginasi cocok.
- Cover depan + biografi belakang (gambar asli), tanpa nomor, di luar hitungan.
- Penomoran isi restart dari 1 (Kata Pengantar = 1).
- DAFTAR ISI MANUAL (nomor halaman diketik, dot leader) -> tidak perlu update field.
- Tiap bab page-break-before (tanpa halaman kosong), heading keep-with-next (tanpa gantung).
"""
import os, zipfile, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from konten import KATA_PENGANTAR, BAB, DAFTAR_PUSTAKA

SRC = "drive_buku/Joki buku pkn/5_6183913965983112173.docx"
OUT = "Joki-tugas-/FINAL/Buku 1 - Nurul Syifa - PKN - FINISH.docx"
PR = json.load(open("build_nurul/page_real.json"))

FONT = "Times New Roman"
TEXT_W_CM = 14.5  # 21 - 3.5(kiri) - 3(kanan)

def extract_covers():
    os.makedirs("build_nurul/img", exist_ok=True)
    z = zipfile.ZipFile(SRC)
    media = sorted([n for n in z.namelist() if n.startswith("word/media/image")])
    paths = []
    for i, m in enumerate(media, 1):
        p = "build_nurul/img/cv%d%s" % (i, os.path.splitext(m)[1])
        open(p, "wb").write(z.read(m)); paths.append(p)
    return paths

def set_font(run, size=12, bold=False, italic=False):
    run.font.name = FONT; run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    rpr = run._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    rf.set(qn('w:eastAsia'), FONT)

def field(paragraph, instr):
    # PAGE field lengkap: begin -> instrText -> separate -> hasil-cache -> end
    r1 = paragraph.add_run(); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'begin'); r1._r.append(e); set_font(r1, 11)
    r2 = paragraph.add_run(); it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' %s ' % instr; r2._r.append(it); set_font(r2, 11)
    r3 = paragraph.add_run(); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'separate'); r3._r.append(e); set_font(r3, 11)
    r4 = paragraph.add_run('1'); set_font(r4, 11)  # hasil cache (akan di-update Word)
    r5 = paragraph.add_run(); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end'); r5._r.append(e); set_font(r5, 11)

def enable_update_fields(doc):
    # paksa Word menghitung ulang semua field saat dokumen dibuka
    settings = doc.settings.element
    if settings.find(qn('w:updateFields')) is None:
        u = OxmlElement('w:updateFields'); u.set(qn('w:val'), 'true'); settings.append(u)

def pgnum_start(section, start=1):
    sectPr = section._sectPr
    pg = sectPr.find(qn('w:pgNumType'))
    if pg is None:
        pg = OxmlElement('w:pgNumType'); sectPr.append(pg)
    pg.set(qn('w:start'), str(start))

def para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, after=12, before=0, italic=False, bold=False):
    p = doc.add_paragraph(); pf = p.paragraph_format
    p.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.75
    pf.space_after = Pt(after); pf.space_before = Pt(before)
    r = p.add_run(text); set_font(r, size, bold=bold, italic=italic)
    return p

def heading(doc, text, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18, page_break=False, keep_next=True):
    p = doc.add_paragraph(); pf = p.paragraph_format
    p.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.75
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.keep_with_next = keep_next
    if page_break:
        pf.page_break_before = True
    r = p.add_run(text); set_font(r, size, bold=True)
    return p

def full_cover(doc, img, first=False):
    sec = doc.sections[0] if first else doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Cm(0)
    sec.header_distance = Cm(0); sec.footer_distance = Cm(0)
    sec.footer.is_linked_to_previous = False
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(img, width=Cm(21), height=Cm(29.7))
    return sec

def toc_entry(doc, label, page):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
    pf.space_after = Pt(10)
    pf.tab_stops.add_tab_stop(Cm(TEXT_W_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(label + "\t" + str(page)); set_font(r, 12)

def build():
    front, back = extract_covers()
    doc = Document()
    # Normal style default
    st = doc.styles['Normal']; st.font.name = FONT; st.font.size = Pt(12)

    # ---- COVER DEPAN ----
    full_cover(doc, front, first=True)

    # ---- ISI ----
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width = Cm(21); body.page_height = Cm(29.7)
    body.top_margin = Cm(3); body.bottom_margin = Cm(2.5)
    body.left_margin = Cm(3.5); body.right_margin = Cm(3)
    pgnum_start(body, 1)
    body.footer.is_linked_to_previous = False
    fp = body.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field(fp, "PAGE")

    # KATA PENGANTAR
    heading(doc, "KATA PENGANTAR", page_break=False)
    for t in KATA_PENGANTAR:
        para(doc, t)

    # DAFTAR ISI (manual, nomor terbukti benar)
    heading(doc, "DAFTAR ISI", page_break=True)
    toc_entry(doc, "KATA PENGANTAR", PR['kp'])
    for b in BAB:
        toc_entry(doc, "BAB %s  %s" % (b['no'], b['judul']), PR['bab'+b['no']])
    toc_entry(doc, "DAFTAR PUSTAKA", PR['dp'])

    # BAB 1-8
    for b in BAB:
        heading(doc, "BAB %s" % b['no'], size=16, page_break=True, after=4)
        heading(doc, b['judul'], size=16, page_break=False, before=0, after=20)
        for k, t in b['isi']:
            if k == 'h2':
                heading(doc, t, size=13, align=WD_ALIGN_PARAGRAPH.LEFT, before=14, after=6, page_break=False)
            else:
                para(doc, t)

    # DAFTAR PUSTAKA
    heading(doc, "DAFTAR PUSTAKA", size=16, page_break=True, after=18)
    for r in DAFTAR_PUSTAKA:
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.75
        pf.space_after = Pt(8); pf.left_indent = Cm(1); pf.first_line_indent = Cm(-1)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(r); set_font(run, 12)

    # ---- COVER BELAKANG ----
    full_cover(doc, back, first=False)

    enable_update_fields(doc)
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("SAVED:", OUT, os.path.getsize(OUT), "bytes")

if __name__ == "__main__":
    build()
