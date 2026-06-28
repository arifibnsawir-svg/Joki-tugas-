# -*- coding: utf-8 -*-
import re, zipfile, sys
from pypdf import PdfReader
import docx
sys.path.insert(0, "Joki-tugas-/_generator/buku6")
import konten

PDF = "Joki-tugas-/FINAL/Buku 6 - Revalina Damayanti - PKN - FINISH.pdf"
DOC = "Joki-tugas-/FINAL/Buku 6 - Revalina Damayanti - PKN - FINISH.docx"
WAJIB = [(1, "HAKIKAT PENDIDIKAN KEWARGANEGARAAN"), (2, "IDENTITAS NASIONAL"), (3, "INTEGRASI NASIONAL"),
(4, "NEGARA DAN KONSTITUSI"), (5, "HAK DAN KEWAJIBAN WARGA NEGARA"), (6, "PENEGAKAN HUKUM YANG BERKEADILAN"),
(7, "GEOPOLITIK DAN GEOSTRATEGI"), (8, "ANTI KORUPSI")]
res = []
def add(n, ok, d=""): res.append((n, ok, d))

r = PdfReader(PDF); n = len(r.pages); pages = [(p.extract_text() or "") for p in r.pages]; full = "\n".join(pages)
norm = " ".join(full.split())
content = n - 2
add("Min. 60 halaman isi (di luar cover)", content >= 60, f"{content} halaman isi (total {n} - 2 cover)")
add("Maks. 80 halaman isi (aturan baru)", content <= 80, f"{content} halaman isi")

# COVER DEPAN: judul lengkap, nama penulis, TANPA kata 'minibook'/'modul pembelajaran'
cover0 = " ".join(pages[0].split())
front_ok = ("PENDIDIKAN KEWARGANEGARAAN" in cover0 and "DI PERGURUAN TINGGI" in cover0 and "Revalina Damayanti" in cover0)
add("Cover depan: judul persis + nama penulis", front_ok, cover0[:90])
no_minibook = ("minibook" not in cover0.lower()) and ("mini book" not in cover0.lower()) and ("modul pembelajaran" not in cover0.lower())
add("Cover depan TANPA kata 'minibook'/'modul'", no_minibook, "tidak ada kata minibook/modul di cover depan")

# COVER BELAKANG: seirama (didesain, ada teks) + biografi penulis
cover_last = " ".join(pages[-1].split())
back_ok = ("BIOGRAFI PENULIS" in cover_last) and ("Revalina Damayanti" in cover_last) and ("Bimbingan dan Konseling" in cover_last)
add("Cover belakang: biografi penulis", back_ok, "ada heading BIOGRAFI PENULIS + nama + prodi")

# 8 bab urut + judul persis. Deteksi via paragraf pertama tiap bab (hanya ada di ISI, bukan di
# Daftar Isi yang ter-uppercase oleh text-transform), lalu cek judul bab uppercase ada di halaman itu.
def joinp(i): return " ".join(pages[i].split())
starts = {}
judul_ok = True
for no, jd in WAJIB:
    b = next(bb for bb in konten.BAB if bb['no'] == str(no))
    first_p = next((blk[1] for blk in b['isi'] if blk[0] == 'p'), "")
    probe = " ".join(first_p.split())[:55]
    for i in range(1, n - 1):
        if probe and probe in joinp(i):
            starts[no] = i
            break
    if no not in starts or jd not in joinp(starts[no]):
        judul_ok = False
order = True; last = -1; det = []
for no, jd in WAJIB:
    fp_ = starts.get(no)
    if fp_ is None: order = False
    else:
        if fp_ <= last: order = False
        last = fp_
    det.append(f"B{no}={fp_}")
add("8 bab sesuai silabus & urut", order and len(starts) == 8, " ".join(det))
add("Judul bab PERSIS silabus", judul_ok and len(starts) == 8, "semua judul bab cocok kata demi kata")

def fp(label):
    for i, t in enumerate(pages):
        if " ".join(t.split()).upper().startswith(label): return i
    return None
kp, di, dp = fp("KATA PENGANTAR"), fp("DAFTAR ISI"), fp("DAFTAR PUSTAKA")
add("Sistematika urut (KP->DI->Bab->DP)", None not in (kp, di, dp) and kp < di < dp, f"KP=hal{kp} DI=hal{di} DP=hal{dp}")

bab1 = starts.get(1, di + 1 if di else 1)
toc_text = "".join(pages[di:bab1]) if di is not None else ""
dots = len(re.findall(r"\.{3,}", toc_text))
bab_in_toc = len(re.findall(r"BAB \d", toc_text))
add("Daftar Isi nomor rata kanan + garis (tanpa dot-leader)", dots == 0 and bab_in_toc >= 8, f"dot-leader={dots} (harus 0), entri BAB di TOC={bab_in_toc}")

add("Daftar Pustaka + tautan", len(re.findall(r"https?://", full)) >= 8, f"{len(re.findall(r'https?://', full))} tautan")
# Vancouver: bernomor '1.' di awal entri + 'Tersedia dari:' + '[Internet]' (beda dari IEEE '[n]'/'Tersedia pada:')
vanc = ("Tersedia dari:" in full) and ("[Internet]" in full) and (re.search(r"\b1\.\s", norm) is not None) and ("Tersedia pada:" not in full)
add("Daftar Pustaka gaya Vancouver", vanc, "bernomor '1.', '[Internet]', 'Tersedia dari:'")

blank = [i + 1 for i in range(1, n - 1) if len(pages[i].strip()) < 5]
add("Tanpa halaman kosong di isi", len(blank) == 0, f"{blank if blank else 'tidak ada'}")
bad = sum(full.count(c) for c in ["\u2014", "\u2013", "\u201c", "\u201d", "\u2018", "\u2019"])
emoji = [c for c in full if ord(c) > 0x2190]
add("Humanizer (no em-dash/kutip keriting/emoji)", bad == 0 and len(emoji) == 0, f"em/en-dash+kutip:{bad}, emoji:{len(emoji)}")

# DNA pembeda Buku 6
nakar = len(re.findall(r"Akar Masalah", full))
add("DNA: kotak 'Akar Masalah'", nakar >= 8, f"{nakar} kotak Akar Masalah terdeteksi")
nsol = len(re.findall(r"Langkah Solusi", full))
add("DNA: kotak 'Langkah Solusi'", nsol >= 8, f"{nsol} kotak Langkah Solusi terdeteksi")

# DOCX
d = docx.Document(DOC)
add("DOCX: ada (editable)", True, f"{len(d.sections)} section, {len(d.inline_shapes)} gambar, {len(d.tables)} kotak")
babs = [p.text.strip() for p in d.paragraphs if re.match(r"^BAB \d$", p.text.strip())]
add("DOCX: 8 bab", len(babs) == 8, " ".join(babs))
zf = zipfile.ZipFile(DOC)
hdr = any('PAGE' in zf.read(x).decode('utf-8', 'ignore') for x in zf.namelist() if 'header' in x)
ftr = any('PAGE' in zf.read(x).decode('utf-8', 'ignore') for x in zf.namelist() if 'footer' in x)
add("DOCX: nomor halaman di HEADER (pojok kanan atas)", hdr and not ftr, f"header PAGE={hdr}, footer PAGE={ftr}")
add("DOCX: ada kotak (Akar Masalah + Langkah Solusi)", len(d.tables) >= 16, f"{len(d.tables)} kotak (target >=16)")
add("DOCX: cover depan & belakang (gambar)", len(d.inline_shapes) >= 2, f"{len(d.inline_shapes)} gambar cover")

print("=" * 66); print("QA BUKU 6 - REVALINA DAMAYANTI  vs  INSTRUKSI DOSEN + DNA"); print("=" * 66)
allp = True
for nm, ok, d2 in res:
    print(f"[{'PASS' if ok else 'FAIL'}] {nm}")
    if d2: print(f"        -> {d2}")
    if not ok: allp = False
print("=" * 66); print("HASIL:", "SEMUA PASS" if allp else "ADA FAIL")
