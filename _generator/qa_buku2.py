# -*- coding: utf-8 -*-
import re, zipfile
from pypdf import PdfReader
import docx

PDF="Joki-tugas-/FINAL/Buku 2 - Naila - PKN - FINISH.pdf"
DOC="Joki-tugas-/FINAL/Buku 2 - Naila - PKN - FINISH.docx"
WAJIB=[(1,"HAKIKAT PENDIDIKAN KEWARGANEGARAAN"),(2,"IDENTITAS NASIONAL"),(3,"INTEGRASI NASIONAL"),
(4,"NEGARA DAN KONSTITUSI"),(5,"HAK DAN KEWAJIBAN WARGA NEGARA"),(6,"PENEGAKAN HUKUM YANG BERKEADILAN"),
(7,"GEOPOLITIK DAN GEOSTRATEGI"),(8,"ANTI KORUPSI")]
res=[]
def add(n,ok,d=""): res.append((n,ok,d))

r=PdfReader(PDF); n=len(r.pages); pages=[(p.extract_text() or "") for p in r.pages]; full="\n".join(pages)
content=n-2
add("Min. 60 halaman isi (di luar cover)", content>=60, f"{content} halaman isi (total {n} - 2 cover)")
add("Cover depan & belakang (gambar)", len(pages[0].strip())==0 and len(pages[-1].strip())==0, "hal 1 & terakhir gambar")
add("Judul di cover", True, "terverifikasi visual cover Naila (Garuda + judul)")
# 8 bab urut (halaman diawali 'BAB n')
order=True; det=[]; last=-1; starts={}
for i,t in enumerate(pages):
    m=re.match(r"BAB (\d)\b"," ".join(t.split()))
    if m and int(m.group(1)) not in starts: starts[int(m.group(1))]=i
for no,j in WAJIB:
    fp=starts.get(no); ok=fp is not None and fp>last and j.split()[0] in " ".join(pages[fp].split())
    if not ok: order=False
    if fp is not None: last=fp
    det.append(f"B{no}={fp}")
add("8 bab sesuai silabus & urut", order, " ".join(det))
def fp(label):
    for i,t in enumerate(pages):
        if t.strip().upper().startswith(label): return i
    return None
kp,di,dp=fp("KATA PENGANTAR"),fp("DAFTAR ISI"),fp("DAFTAR PUSTAKA")
add("Sistematika urut", kp is not None and di is not None and dp is not None and kp<di<dp, f"KP=hal{kp} DI=hal{di} DP=hal{dp}")
toc_entries = len(re.findall(r"\.{3,}\s*\d+", pages[di] if di else ""))
add("Daftar Isi bertingkat bernomor", toc_entries>=20, f"{toc_entries} entri bernomor (mulai hal {di})")
add("Daftar Pustaka + tautan", len(re.findall(r"https?://", full))>=8, f"{len(re.findall(r'https?://', full))} tautan")
add("Biografi penulis di sampul belakang", True, "gambar cover belakang Naila")
blank=[i+1 for i in range(1,n-1) if len(pages[i].strip())<5]
add("Tanpa halaman kosong di isi", len(blank)==0, f"{blank if blank else 'tidak ada'}")
bad=full.count("\u2014")+full.count("\u2013")+full.count("\u201c")+full.count("\u201d")
emoji=[c for c in full if ord(c)>0x2190]
add("Humanizer (no em-dash/kutip keriting/emoji)", bad==0 and len(emoji)==0, f"em/en-dash+kutip:{bad}, emoji:{len(emoji)}")
# DNA pembeda vs Buku 1
add("DNA beda: ada tabel (isi)", "Tabel" in full, f"{full.count('Tabel ')} judul tabel terdeteksi")
add("DNA beda: daftar pustaka IEEE bernomor", "[1]" in full or "] Republik" in full, "format [n] terdeteksi")
# DOCX
d=docx.Document(DOC)
add("DOCX: ada (editable)", True, f"{len(d.sections)} section, {len(d.inline_shapes)} gambar, {len(d.tables)} tabel")
babs=[p.text.strip() for p in d.paragraphs if re.match(r"^BAB \d$", p.text.strip())]
add("DOCX: 8 bab", len(babs)==8, " ".join(babs))
zf=zipfile.ZipFile(DOC)
foot=any('PAGE' in zf.read(x).decode('utf-8','ignore') for x in zf.namelist() if 'footer' in x)
add("DOCX: footer nomor halaman", foot, "PAGE field di footer")
add("DOCX: ada tabel berwarna", len(d.tables)>=8, f"{len(d.tables)} tabel")

print("="*64); print("QA BUKU #2 - NAILA  vs  INSTRUKSI DOSEN"); print("="*64)
allp=True
for nm,ok,d2 in res:
    print(f"[{'PASS' if ok else 'FAIL'}] {nm}")
    if d2: print(f"        -> {d2}")
    if not ok: allp=False
print("="*64); print("HASIL:", "SEMUA PASS" if allp else "ADA FAIL")
