# -*- coding: utf-8 -*-
"""DOCX Buku #5 (Nurjali Sangadji) - mirror layout PDF (A4, TNR 12, spasi 1.9).
DNA: kotak Poin Inti (bernomor) + kotak Istilah Kunci (indigo), nomor halaman BAWAH-TENGAH
(footer dgn garis atas), daftar isi MANUAL banded (tanpa leader), pustaka Chicago.
"""
import os, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from konten import KATA_PENGANTAR, BAB, DAFTAR_PUSTAKA

OUT = "Joki-tugas-/FINAL/Buku 5 - Nurjali Sangadji - PKN - FINISH.docx"
IMGDIR = "Joki-tugas-/_generator/buku5/build_nurjali/img"
PR = json.load(open("Joki-tugas-/_generator/buku5/build_nurjali/page_real.json"))
FONT = "Times New Roman"
TEXT_W_CM = 14.5
LS = 1.9
INDIGO = RGBColor(0x2B, 0x2F, 0x5E)

def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT; run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'), FONT)

def ls_set(pf, val=LS):
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = val

def ppr_flag(p, tag):
    pPr = p._p.get_or_add_pPr(); e = OxmlElement('w:' + tag); e.set(qn('w:val'), 'true'); pPr.append(e)

def widow(p): ppr_flag(p, 'widowControl')

def shade_par(p, fill):
    pPr = p._p.get_or_add_pPr(); sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill); pPr.append(sh)

def border_top(p, color="8A8FB5", sz="4"):
    pPr = p._p.get_or_add_pPr(); pbd = OxmlElement('w:pBdr'); top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), sz); top.set(qn('w:space'), '4'); top.set(qn('w:color'), color)
    pbd.append(top); pPr.append(pbd)

def page_field(paragraph):
    r = paragraph.add_run(); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'begin'); r._r.append(e); set_font(r, 11)
    r = paragraph.add_run(); it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '; r._r.append(it); set_font(r, 11)
    r = paragraph.add_run(); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'separate'); r._r.append(e); set_font(r, 11)
    r = paragraph.add_run('1'); set_font(r, 11)
    r = paragraph.add_run(); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end'); r._r.append(e); set_font(r, 11)

def enable_update_fields(doc):
    s = doc.settings.element
    if s.find(qn('w:updateFields')) is None:
        u = OxmlElement('w:updateFields'); u.set(qn('w:val'), 'true'); s.append(u)

def pgnum_start(section, start=1):
    sectPr = section._sectPr; pg = sectPr.find(qn('w:pgNumType'))
    if pg is None: pg = OxmlElement('w:pgNumType'); sectPr.append(pg)
    pg.set(qn('w:start'), str(start))

def para(doc, text, size=12, after=10, before=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(); pf = p.paragraph_format; p.alignment = align; ls_set(pf)
    pf.space_after = Pt(after); pf.space_before = Pt(before); widow(p)
    r = p.add_run(text); set_font(r, size); return p

def heading(doc, text, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18, page_break=False, keep=True, color=INDIGO):
    p = doc.add_paragraph(); pf = p.paragraph_format; p.alignment = align; ls_set(pf)
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.keep_with_next = keep
    widow(p); ppr_flag(p, 'keepLines')
    if page_break: pf.page_break_before = True
    r = p.add_run(text); set_font(r, size, bold=True, color=color); return p

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill); tcPr.append(sh)

def box_borders(table, left='4A4F93', rest='A9AEDE', dash=False):
    tbl = table._tbl; tblPr = tbl.tblPr; borders = OxmlElement('w:tblBorders')
    val = 'dashed' if dash else 'single'
    spec = {'top': ('6', rest), 'bottom': ('6', rest), 'right': ('6', rest), 'left': ('24', left), 'insideH': ('6', rest), 'insideV': ('6', rest)}
    for edge, (sz, col) in spec.items():
        e = OxmlElement('w:' + edge); e.set(qn('w:val'), val if edge != 'left' else 'single'); e.set(qn('w:sz'), sz); e.set(qn('w:color'), col); borders.append(e)
    tblPr.append(borders)

def add_poin(doc, p):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    box_borders(t)
    cell = t.rows[0].cells[0]; shade_cell(cell, 'EEF0F8')
    p0 = cell.paragraphs[0]; p0.paragraph_format.space_after = Pt(5); ls_set(p0.paragraph_format, 1.4); widow(p0)
    r = p0.add_run(p['judul']); set_font(r, 11.5, bold=True, color=INDIGO)
    for i, it in enumerate(p['items'], 1):
        pi = cell.add_paragraph(); pf = pi.paragraph_format; pf.left_indent = Cm(0.8); pf.first_line_indent = Cm(-0.5)
        ls_set(pf, 1.4); pf.space_after = Pt(3); pi.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; widow(pi)
        rr = pi.add_run("%d. %s" % (i, it)); set_font(rr, 11)
    rpr = t.rows[0]._tr.get_or_add_trPr(); cs = OxmlElement('w:cantSplit'); cs.set(qn('w:val'), 'true'); rpr.append(cs)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_istilah(doc, t0):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    box_borders(t, left='8A8FC4', rest='8A8FC4', dash=True)
    cell = t.rows[0].cells[0]; shade_cell(cell, 'F5F6FB')
    p0 = cell.paragraphs[0]; p0.paragraph_format.space_after = Pt(5); ls_set(p0.paragraph_format, 1.4); widow(p0)
    r = p0.add_run(t0['judul']); set_font(r, 11.5, bold=True, color=RGBColor(0x3A, 0x3F, 0x80))
    for term, arti in t0['pairs']:
        pi = cell.add_paragraph(); pf = pi.paragraph_format; ls_set(pf, 1.45); pf.space_after = Pt(5); pi.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; widow(pi)
        rb = pi.add_run(term + ": "); set_font(rb, 11, bold=True, color=INDIGO); ra = pi.add_run(arti); set_font(ra, 11)
    rpr = t.rows[0]._tr.get_or_add_trPr(); cs = OxmlElement('w:cantSplit'); cs.set(qn('w:val'), 'true'); rpr.append(cs)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def full_cover(doc, img, first=False):
    sec = doc.sections[0] if first else doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Cm(0)
    sec.header_distance = Cm(0); sec.footer_distance = Cm(0); sec.footer.is_linked_to_previous = False
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(img, width=Cm(21), height=Cm(29.7))

def toc_line(doc, label, page, bab=True):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.2
    pf.space_after = Pt(3 if bab else 2); pf.space_before = Pt(5 if bab else 0)
    if not bab: pf.left_indent = Cm(0.9)
    pf.tab_stops.add_tab_stop(Cm(TEXT_W_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    if bab:
        shade_par(p, 'E9EBF6')
        pf.left_indent = Cm(0.2)
    r = p.add_run(label + "\t" + str(page)); set_font(r, 12, bold=bab, color=INDIGO if bab else None)

def build():
    front = os.path.join(IMGDIR, "cv1.jpg"); back = os.path.join(IMGDIR, "cv2.jpg")
    doc = Document(); doc.styles['Normal'].font.name = FONT; doc.styles['Normal'].font.size = Pt(12)
    full_cover(doc, front, first=True)
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width = Cm(21); body.page_height = Cm(29.7)
    body.top_margin = Cm(3); body.bottom_margin = Cm(2.5); body.left_margin = Cm(3.5); body.right_margin = Cm(3)
    pgnum_start(body, 1); body.footer.is_linked_to_previous = False
    fp = body.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER  # BAWAH-TENGAH
    border_top(fp); page_field(fp)

    heading(doc, "KATA PENGANTAR")
    for t in KATA_PENGANTAR: para(doc, t)

    heading(doc, "DAFTAR ISI", page_break=True)
    toc_line(doc, "Kata Pengantar", PR['kp'], bab=True)
    for b in BAB:
        toc_line(doc, "BAB %s  %s" % (b['no'], b['judul'].title()), PR['bab' + b['no']], bab=True)
        k = 0
        for blk in b['isi']:
            if blk[0] == 'h2':
                k += 1
                toc_line(doc, blk[1], PR.get("s%s_%d" % (b['no'], k), ""), bab=False)
    toc_line(doc, "Daftar Pustaka", PR['dp'], bab=True)

    for b in BAB:
        heading(doc, "BAB %s" % b['no'], size=16, page_break=True, after=4)
        heading(doc, b['judul'], size=16, after=18)
        for blk in b['isi']:
            if blk[0] == 'h2':
                heading(doc, blk[1], size=13, align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=6)
            elif blk[0] == 'poin':
                add_poin(doc, blk[1])
            elif blk[0] == 'istilah':
                add_istilah(doc, blk[1])
            else:
                para(doc, blk[1])

    heading(doc, "DAFTAR PUSTAKA", size=16, page_break=True, after=18)
    for pre, title, post in DAFTAR_PUSTAKA:
        p = doc.add_paragraph(); pf = p.paragraph_format; ls_set(pf)
        pf.space_after = Pt(8); pf.left_indent = Cm(1.0); pf.first_line_indent = Cm(-1.0); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        widow(p)
        r1 = p.add_run(pre); set_font(r1, 12)
        r2 = p.add_run(title); set_font(r2, 12, italic=True)
        r3 = p.add_run(post); set_font(r3, 12)

    full_cover(doc, back, first=False)
    enable_update_fields(doc)
    os.makedirs(os.path.dirname(OUT), exist_ok=True); doc.save(OUT)
    print("SAVED:", OUT, os.path.getsize(OUT), "bytes")

if __name__ == "__main__":
    build()
