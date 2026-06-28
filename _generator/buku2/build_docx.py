# -*- coding: utf-8 -*-
"""DOCX Buku #2 (Naila) - mirror layout PDF (A4, TNR 12, spasi 2.0).
DNA: tabel berwarna, nomor halaman POJOK KANAN BAWAH, daftar isi MANUAL bertingkat.
"""
import os, zipfile, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from konten import KATA_PENGANTAR, BAB, DAFTAR_PUSTAKA

SRC = "Joki-tugas-/SUMBER_ASLI (jangan dikumpulkan)/MINI_BOOK_PKN_NAILA_FINAL_V10_COVER_INTEGRATED finish.docx"
OUT = "Joki-tugas-/FINAL/Buku 2 - Naila - PKN - FINISH.docx"
PR = json.load(open("Joki-tugas-/_generator/buku2/build_naila/page_real.json"))
FONT = "Times New Roman"
TEXT_W_CM = 14.5

def covers_extract():
    os.makedirs("Joki-tugas-/_generator/buku2/build_naila/img", exist_ok=True)
    z = zipfile.ZipFile(SRC)
    media = sorted([n for n in z.namelist() if n.startswith("word/media/image")])
    paths = []
    for i, m in enumerate(media, 1):
        p = "Joki-tugas-/_generator/buku2/build_naila/img/cv%d%s" % (i, os.path.splitext(m)[1]); open(p,"wb").write(z.read(m)); paths.append(p)
    return paths

def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT; run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'), FONT)

def ls_double(pf):
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.6

def page_field(paragraph):
    for typ in ('begin',):
        r=paragraph.add_run(); e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),typ); r._r.append(e); set_font(r,11)
    r=paragraph.add_run(); it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text=' PAGE '; r._r.append(it); set_font(r,11)
    r=paragraph.add_run(); e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'separate'); r._r.append(e); set_font(r,11)
    r=paragraph.add_run('1'); set_font(r,11)
    r=paragraph.add_run(); e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); r._r.append(e); set_font(r,11)

def enable_update_fields(doc):
    s=doc.settings.element
    if s.find(qn('w:updateFields')) is None:
        u=OxmlElement('w:updateFields'); u.set(qn('w:val'),'true'); s.append(u)

def pgnum_start(section, start=1):
    sectPr=section._sectPr; pg=sectPr.find(qn('w:pgNumType'))
    if pg is None: pg=OxmlElement('w:pgNumType'); sectPr.append(pg)
    pg.set(qn('w:start'), str(start))

def ppr_flag(p, tag):
    pPr=p._p.get_or_add_pPr(); e=OxmlElement('w:'+tag); e.set(qn('w:val'),'true'); pPr.append(e)

def widow(p):
    ppr_flag(p, 'widowControl')

def para(doc, text, size=12, after=10, before=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p=doc.add_paragraph(); pf=p.paragraph_format; p.alignment=align; ls_double(pf)
    pf.space_after=Pt(after); pf.space_before=Pt(before)
    widow(p)
    r=p.add_run(text); set_font(r,size); return p

def heading(doc, text, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18, page_break=False, keep=True):
    p=doc.add_paragraph(); pf=p.paragraph_format; p.alignment=align; ls_double(pf)
    pf.space_before=Pt(before); pf.space_after=Pt(after); pf.keep_with_next=keep
    widow(p); ppr_flag(p, 'keepLines')  # heading tidak terbelah & tidak menggantung
    if page_break: pf.page_break_before=True
    r=p.add_run(text); set_font(r,size,bold=True); return p

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),fill); tcPr.append(sh)

def set_borders(table):
    tbl=table._tbl; tblPr=tbl.tblPr; borders=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e=OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),'8AA0C8'); borders.append(e)
    tblPr.append(borders)

def add_table(doc, tb):
    cap=doc.add_paragraph(); cap.paragraph_format.space_after=Pt(3); cap.paragraph_format.space_before=Pt(6)
    cap.paragraph_format.keep_with_next=True  # caption nempel ke tabel
    widow(cap)
    rc=cap.add_run(tb['judul']); set_font(rc,10.5,italic=True,color=RGBColor(0x1F,0x38,0x64))
    t=doc.add_table(rows=1, cols=len(tb['head'])); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    set_borders(t)
    hdr=t.rows[0].cells
    for j,h in enumerate(tb['head']):
        shade(hdr[j],'D9E2F3'); pr=hdr[j].paragraphs[0]; pr.paragraph_format.space_after=Pt(2)
        rr=pr.add_run(h); set_font(rr,11,bold=True,color=RGBColor(0x1F,0x38,0x64))
    # header row diulang bila tabel terpotong
    trPr=t.rows[0]._tr.get_or_add_trPr(); th=OxmlElement('w:tblHeader'); th.set(qn('w:val'),'true'); trPr.append(th)
    for ri,row in enumerate(tb['rows']):
        cells=t.add_row().cells
        for j,c in enumerate(row):
            if ri%2==1: shade(cells[j],'F2F6FC')
            pr=cells[j].paragraphs[0]; pr.paragraph_format.space_after=Pt(2)
            rr=pr.add_run(c); set_font(rr,11)
    # tiap baris tidak terbelah antar-halaman
    for row in t.rows:
        rpr=row._tr.get_or_add_trPr(); cs=OxmlElement('w:cantSplit'); cs.set(qn('w:val'),'true'); rpr.append(cs)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

def full_cover(doc, img, first=False):
    sec=doc.sections[0] if first else doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    sec.left_margin=sec.right_margin=sec.top_margin=sec.bottom_margin=Cm(0)
    sec.header_distance=Cm(0); sec.footer_distance=Cm(0); sec.footer.is_linked_to_previous=False
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
    p.add_run().add_picture(img, width=Cm(21), height=Cm(29.7))

def toc_line(doc, label, page, lvl1=True):
    p=doc.add_paragraph(); pf=p.paragraph_format
    pf.line_spacing_rule=WD_LINE_SPACING.MULTIPLE; pf.line_spacing=1.3
    pf.space_after=Pt(6 if lvl1 else 3)
    if not lvl1: pf.left_indent=Cm(1.0)
    pf.tab_stops.add_tab_stop(Cm(TEXT_W_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r=p.add_run(label+"\t"+str(page)); set_font(r,12,bold=lvl1)

def build():
    front, back = covers_extract()
    doc=Document(); doc.styles['Normal'].font.name=FONT; doc.styles['Normal'].font.size=Pt(12)
    full_cover(doc, front, first=True)
    body=doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width=Cm(21); body.page_height=Cm(29.7)
    body.top_margin=Cm(3); body.bottom_margin=Cm(2.5); body.left_margin=Cm(3.5); body.right_margin=Cm(3)
    pgnum_start(body,1); body.footer.is_linked_to_previous=False
    fp=body.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT  # POJOK KANAN BAWAH
    page_field(fp)

    # KATA PENGANTAR
    heading(doc, "KATA PENGANTAR")
    for t in KATA_PENGANTAR: para(doc, t)

    # DAFTAR ISI bertingkat (manual)
    heading(doc, "DAFTAR ISI", page_break=True)
    toc_line(doc, "KATA PENGANTAR", PR['kp'], lvl1=True)
    for b in BAB:
        toc_line(doc, "BAB %s  %s" % (b['no'], b['judul']), PR['bab'+b['no']], lvl1=True)
        k=0
        for blk in b['isi']:
            if blk[0]=='h2':
                k+=1
                toc_line(doc, blk[1], PR.get("s%s_%d"%(b['no'],k),""), lvl1=False)
    toc_line(doc, "DAFTAR PUSTAKA", PR['dp'], lvl1=True)

    # BAB
    for b in BAB:
        heading(doc, "BAB %s" % b['no'], size=16, page_break=True, after=4)
        heading(doc, b['judul'], size=16, after=18)
        for blk in b['isi']:
            if blk[0]=='h2':
                heading(doc, blk[1], size=13, align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=6)
            elif blk[0]=='table':
                add_table(doc, blk[1])
            else:
                para(doc, blk[1])

    # DAFTAR PUSTAKA (IEEE bernomor manual)
    heading(doc, "DAFTAR PUSTAKA", size=16, page_break=True, after=18)
    for i, ref in enumerate(DAFTAR_PUSTAKA, 1):
        p=doc.add_paragraph(); pf=p.paragraph_format; ls_double(pf)
        pf.space_after=Pt(8); pf.left_indent=Cm(1.0); pf.first_line_indent=Cm(-1.0); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r=p.add_run("[%d] %s" % (i, ref)); set_font(r,12)

    full_cover(doc, back, first=False)
    enable_update_fields(doc)
    os.makedirs(os.path.dirname(OUT), exist_ok=True); doc.save(OUT)
    print("SAVED:", OUT, os.path.getsize(OUT), "bytes")

if __name__ == "__main__":
    build()
