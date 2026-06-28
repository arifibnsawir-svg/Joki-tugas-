# -*- coding: utf-8 -*-
"""DOCX Buku #3 (Balqis Lejla) - mirror layout PDF (A4, TNR 12, spasi 1.9 seperti PDF).
DNA: poin-poin/bullet + kotak studi kasus (warna hangat), nomor halaman POJOK LUAR berganti
(genap kiri, ganjil kanan) via evenAndOddHeaders, daftar isi MANUAL leader titik (nomor bab
berformat zero-pad), daftar pustaka Harvard. Daftar isi pakai nomor dari scan PDF (acuan=PDF).
"""
import os, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from konten import KATA_PENGANTAR, BAB, DAFTAR_PUSTAKA

OUT = "Joki-tugas-/FINAL/Buku 3 - Balqis Lejla - PKN - FINISH.docx"
IMGDIR = "Joki-tugas-/_generator/buku3/build_balqis/img"
PR = json.load(open("Joki-tugas-/_generator/buku3/build_balqis/page_real.json"))
FONT = "Times New Roman"
TEXT_W_CM = 14.5
LS = 1.9

def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT; run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'), FONT)

def ls_set(pf, val=LS):
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = val

def ppr_flag(p, tag):
    pPr = p._p.get_or_add_pPr(); e = OxmlElement('w:' + tag); e.set(qn('w:val'), 'true'); pPr.append(e)

def widow(p): ppr_flag(p, 'widowControl')

def page_field(paragraph):
    r = paragraph.add_run(); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'begin'); r._r.append(e); set_font(r, 11)
    r = paragraph.add_run(); it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '; r._r.append(it); set_font(r, 11)
    r = paragraph.add_run(); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'separate'); r._r.append(e); set_font(r, 11)
    r = paragraph.add_run('1'); set_font(r, 11)
    r = paragraph.add_run(); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end'); r._r.append(e); set_font(r, 11)

def enable_update_fields(doc):
    s = doc.settings.element
    if s.find(qn('w:updateFields')) is None:
        u = OxmlElement('w:updateFields'); u.set(qn('w:val'), 'true'); s.append(u)

def enable_even_odd(doc):
    s = doc.settings.element
    if s.find(qn('w:evenAndOddHeaders')) is None:
        e = OxmlElement('w:evenAndOddHeaders'); e.set(qn('w:val'), 'true'); s.append(e)

def pgnum_start(section, start=1):
    sectPr = section._sectPr; pg = sectPr.find(qn('w:pgNumType'))
    if pg is None: pg = OxmlElement('w:pgNumType'); sectPr.append(pg)
    pg.set(qn('w:start'), str(start))

def para(doc, text, size=12, after=10, before=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(); pf = p.paragraph_format; p.alignment = align; ls_set(pf)
    pf.space_after = Pt(after); pf.space_before = Pt(before); widow(p)
    r = p.add_run(text); set_font(r, size); return p

def heading(doc, text, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18, page_break=False, keep=True):
    p = doc.add_paragraph(); pf = p.paragraph_format; p.alignment = align; ls_set(pf)
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.keep_with_next = keep
    widow(p); ppr_flag(p, 'keepLines')
    if page_break: pf.page_break_before = True
    r = p.add_run(text); set_font(r, size, bold=True); return p

def add_bullet(doc, text):
    p = doc.add_paragraph(); pf = p.paragraph_format; p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; ls_set(pf)
    pf.space_after = Pt(5); pf.left_indent = Cm(1.0); pf.first_line_indent = Cm(-0.5); widow(p)
    r = p.add_run("\u2022  " + text); set_font(r, 12); return p

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill); tcPr.append(sh)

def kasus_borders(table):
    tbl = table._tbl; tblPr = tbl.tblPr; borders = OxmlElement('w:tblBorders')
    spec = {'top': ('6', 'D2B48C'), 'bottom': ('6', 'D2B48C'), 'right': ('6', 'D2B48C'),
            'left': ('24', 'B8860B'), 'insideH': ('6', 'D2B48C'), 'insideV': ('6', 'D2B48C')}
    for edge, (sz, col) in spec.items():
        e = OxmlElement('w:' + edge); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz); e.set(qn('w:color'), col); borders.append(e)
    tblPr.append(borders)

def add_kasus(doc, k):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    kasus_borders(t)
    cell = t.rows[0].cells[0]; shade(cell, 'FBF6EC')
    # judul
    p0 = cell.paragraphs[0]; p0.paragraph_format.space_after = Pt(5); ls_set(p0.paragraph_format, 1.4); widow(p0)
    r = p0.add_run(k['judul']); set_font(r, 11.5, bold=True, color=RGBColor(0x6B, 0x45, 0x00))
    for par in k['isi']:
        p = cell.add_paragraph(); pf = p.paragraph_format; p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ls_set(pf, 1.5); pf.space_after = Pt(6); widow(p)
        rr = p.add_run(par); set_font(rr, 11)
    pr = cell.add_paragraph(); prf = pr.paragraph_format; pr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ls_set(prf, 1.4); prf.space_after = Pt(2); widow(pr)
    rr = pr.add_run("Refleksi: " + k['refleksi']); set_font(rr, 10.5, italic=True, color=RGBColor(0x5A, 0x3D, 0x00))
    # cantSplit row
    rpr = t.rows[0]._tr.get_or_add_trPr(); cs = OxmlElement('w:cantSplit'); cs.set(qn('w:val'), 'true'); rpr.append(cs)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def full_cover(doc, img, first=False):
    sec = doc.sections[0] if first else doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Cm(0)
    sec.header_distance = Cm(0); sec.footer_distance = Cm(0); sec.footer.is_linked_to_previous = False
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(img, width=Cm(21), height=Cm(29.7))

def toc_line(doc, label, page, bab=True):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.25
    pf.space_after = Pt(6 if bab else 3)
    if not bab: pf.left_indent = Cm(1.2)
    pf.tab_stops.add_tab_stop(Cm(TEXT_W_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(label + "\t" + str(page)); set_font(r, 12, bold=bab)

def build():
    front = os.path.join(IMGDIR, "cv1.png"); back = os.path.join(IMGDIR, "cv2.png")
    doc = Document(); doc.styles['Normal'].font.name = FONT; doc.styles['Normal'].font.size = Pt(12)
    full_cover(doc, front, first=True)
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width = Cm(21); body.page_height = Cm(29.7)
    body.top_margin = Cm(3); body.bottom_margin = Cm(2.5); body.left_margin = Cm(3.5); body.right_margin = Cm(3)
    pgnum_start(body, 1)
    body.footer.is_linked_to_previous = False
    body.even_page_footer.is_linked_to_previous = False
    # nomor halaman pojok LUAR berganti: ganjil (default) kanan, genap kiri
    fo = body.footer.paragraphs[0]; fo.alignment = WD_ALIGN_PARAGRAPH.RIGHT; page_field(fo)
    fe = body.even_page_footer.paragraphs[0]; fe.alignment = WD_ALIGN_PARAGRAPH.LEFT; page_field(fe)

    # KATA PENGANTAR
    heading(doc, "KATA PENGANTAR")
    for t in KATA_PENGANTAR: para(doc, t)

    # DAFTAR ISI (manual, leader titik, nomor bab zero-pad)
    heading(doc, "DAFTAR ISI", page_break=True)
    toc_line(doc, "Kata Pengantar", PR['kp'], bab=True)
    for b in BAB:
        toc_line(doc, "%02d  %s" % (int(b['no']), b['judul'].title()), PR['bab' + b['no']], bab=True)
        k = 0
        for blk in b['isi']:
            if blk[0] == 'h2':
                k += 1
                toc_line(doc, blk[1], PR.get("s%s_%d" % (b['no'], k), ""), bab=False)
    toc_line(doc, "Daftar Pustaka", PR['dp'], bab=True)

    # BAB
    for b in BAB:
        heading(doc, "BAB %s" % b['no'], size=16, page_break=True, after=4)
        heading(doc, b['judul'], size=16, after=18)
        for blk in b['isi']:
            if blk[0] == 'h2':
                heading(doc, blk[1], size=13, align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=6)
            elif blk[0] == 'ul':
                for it in blk[1]: add_bullet(doc, it)
            elif blk[0] == 'kasus':
                add_kasus(doc, blk[1])
            else:
                para(doc, blk[1])

    # DAFTAR PUSTAKA (Harvard, hanging indent, tanpa nomor, judul miring)
    heading(doc, "DAFTAR PUSTAKA", size=16, page_break=True, after=18)
    for pre, title, post in DAFTAR_PUSTAKA:
        p = doc.add_paragraph(); pf = p.paragraph_format; ls_set(pf)
        pf.space_after = Pt(8); pf.left_indent = Cm(1.0); pf.first_line_indent = Cm(-1.0); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        widow(p)
        r1 = p.add_run(pre); set_font(r1, 12)
        r2 = p.add_run(title); set_font(r2, 12, italic=True)
        r3 = p.add_run(post); set_font(r3, 12)

    full_cover(doc, back, first=False)
    enable_even_odd(doc)
    enable_update_fields(doc)
    os.makedirs(os.path.dirname(OUT), exist_ok=True); doc.save(OUT)
    print("SAVED:", OUT, os.path.getsize(OUT), "bytes")

if __name__ == "__main__":
    build()
