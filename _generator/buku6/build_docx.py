# -*- coding: utf-8 -*-
"""DOCX Buku 6 (Revalina Damayanti) - mirror layout PDF.
DNA: tematik problem-solusi, tema maroon/emas, kotak 'Akar Masalah' + 'Langkah Solusi',
nomor halaman POJOK KANAN ATAS (header), daftar isi NOMOR RATA KANAN + GARIS tiap bab,
pustaka Vancouver. Cover depan & belakang dirasterisasi dari PDF final lalu disisipkan full-page.
"""
import os, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from konten import KATA_PENGANTAR, BAB, DAFTAR_PUSTAKA

PDF = "Joki-tugas-/FINAL/Buku 6 - Revalina Damayanti - PKN - FINISH.pdf"
OUT = "Joki-tugas-/FINAL/Buku 6 - Revalina Damayanti - PKN - FINISH.docx"
ASSETDIR = "Joki-tugas-/_generator/buku6/assets"
PR = json.load(open("Joki-tugas-/_generator/buku6/page_real.json"))
FONT = "Times New Roman"
TEXT_W_CM = 14.5
LS = 1.55
MAROON = RGBColor(0x6E, 0x14, 0x23)
GOLD_DARK = RGBColor(0x7A, 0x5B, 0x12)
GOLD_HEX = "C8A24A"
MAROON_HEX = "6E1423"

def render_covers():
    import fitz
    os.makedirs(ASSETDIR, exist_ok=True)
    d = fitz.open(PDF)
    front = os.path.join(ASSETDIR, "cover_front.png")
    back = os.path.join(ASSETDIR, "cover_back.png")
    mat = fitz.Matrix(200 / 72, 200 / 72)
    d[0].get_pixmap(matrix=mat).save(front)
    d[len(d) - 1].get_pixmap(matrix=mat).save(back)
    return front, back

def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT; run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'), FONT)

def ls_set(pf, val=LS):
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = val

def ppr_flag(p, tag):
    pPr = p._p.get_or_add_pPr(); e = OxmlElement('w:' + tag); e.set(qn('w:val'), 'true'); pPr.append(e)

def widow(p): ppr_flag(p, 'widowControl')

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill); tcPr.append(sh)

def border_bottom(p, color=GOLD_HEX, sz="6"):
    pPr = p._p.get_or_add_pPr(); pbd = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), sz); bot.set(qn('w:space'), '3'); bot.set(qn('w:color'), color)
    pbd.append(bot); pPr.append(pbd)

def page_field(paragraph):
    r = paragraph.add_run(); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'begin'); r._r.append(e); set_font(r, 11, bold=True, color=MAROON)
    r = paragraph.add_run(); it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '; r._r.append(it); set_font(r, 11, bold=True, color=MAROON)
    r = paragraph.add_run(); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'separate'); r._r.append(e); set_font(r, 11, bold=True, color=MAROON)
    r = paragraph.add_run('1'); set_font(r, 11, bold=True, color=MAROON)
    r = paragraph.add_run(); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end'); r._r.append(e); set_font(r, 11, bold=True, color=MAROON)

def enable_update_fields(doc):
    s = doc.settings.element
    if s.find(qn('w:updateFields')) is None:
        u = OxmlElement('w:updateFields'); u.set(qn('w:val'), 'true'); s.append(u)

def pgnum_start(section, start=1):
    sectPr = section._sectPr; pg = sectPr.find(qn('w:pgNumType'))
    if pg is None: pg = OxmlElement('w:pgNumType'); sectPr.append(pg)
    pg.set(qn('w:start'), str(start))

def para(doc, text, size=12, after=10, before=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(); pf = p.paragraph_format; p.alignment = align; ls_set(pf)
    pf.space_after = Pt(after); pf.space_before = Pt(before); widow(p)
    r = p.add_run(text); set_font(r, size); return p

def heading(doc, text, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18, page_break=False, keep=True, color=MAROON, bold=True):
    p = doc.add_paragraph(); pf = p.paragraph_format; p.alignment = align; ls_set(pf)
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.keep_with_next = keep
    widow(p); ppr_flag(p, 'keepLines')
    if page_break: pf.page_break_before = True
    r = p.add_run(text); set_font(r, size, bold=bold, color=color); return p

def box_borders(table, left_hex, rest_hex):
    tbl = table._tbl; tblPr = tbl.tblPr; borders = OxmlElement('w:tblBorders')
    spec = {'top': ('6', rest_hex), 'bottom': ('6', rest_hex), 'right': ('6', rest_hex),
            'left': ('24', left_hex), 'insideH': ('6', rest_hex), 'insideV': ('6', rest_hex)}
    for edge, (sz, col) in spec.items():
        e = OxmlElement('w:' + edge); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz); e.set(qn('w:color'), col); borders.append(e)
    tblPr.append(borders)

def add_box(doc, d, fill, left_hex, rest_hex, title_color, bullet="\u2022 "):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    box_borders(t, left_hex, rest_hex)
    cell = t.rows[0].cells[0]; shade_cell(cell, fill)
    p0 = cell.paragraphs[0]; p0.paragraph_format.space_after = Pt(5); ls_set(p0.paragraph_format, 1.3); widow(p0)
    r = p0.add_run(d['judul']); set_font(r, 12, bold=True, color=title_color)
    for it in d['items']:
        pi = cell.add_paragraph(); pf = pi.paragraph_format; pf.left_indent = Cm(0.8); pf.first_line_indent = Cm(-0.45)
        ls_set(pf, 1.35); pf.space_after = Pt(4); pi.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; widow(pi)
        rr = pi.add_run(bullet + it); set_font(rr, 11)
    rpr = t.rows[0]._tr.get_or_add_trPr(); cs = OxmlElement('w:cantSplit'); cs.set(qn('w:val'), 'true'); rpr.append(cs)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def full_cover(doc, img, first=False):
    sec = doc.sections[0] if first else doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Cm(0)
    sec.header_distance = Cm(0); sec.footer_distance = Cm(0)
    sec.header.is_linked_to_previous = False  # cover tanpa nomor halaman
    sec.footer.is_linked_to_previous = False
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
    p.add_run().add_picture(img, width=Cm(21), height=Cm(29.7))

def toc_line(doc, label, page, bab=True):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.15
    pf.space_after = Pt(2); pf.space_before = Pt(6 if bab else 1)
    pf.tab_stops.add_tab_stop(Cm(TEXT_W_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    if bab:
        border_bottom(p, GOLD_HEX, "6"); pf.keep_with_next = True
        r = p.add_run(label + "\t" + str(page)); set_font(r, 11.5, bold=True, color=MAROON)
    else:
        pf.left_indent = Cm(0.9)
        r = p.add_run(label + "\t"); set_font(r, 11)
        r2 = p.add_run(str(page)); set_font(r2, 11, color=GOLD_DARK)

def build():
    front, back = render_covers()
    doc = Document(); doc.styles['Normal'].font.name = FONT; doc.styles['Normal'].font.size = Pt(12)
    full_cover(doc, front, first=True)
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width = Cm(21); body.page_height = Cm(29.7)
    body.top_margin = Cm(3); body.bottom_margin = Cm(2.5); body.left_margin = Cm(3.5); body.right_margin = Cm(3)
    pgnum_start(body, 1)
    body.header.is_linked_to_previous = False
    hp = body.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # POJOK KANAN ATAS
    page_field(hp)
    body.footer.is_linked_to_previous = False  # footer kosong

    heading(doc, "KATA PENGANTAR")
    for t in KATA_PENGANTAR: para(doc, t)

    heading(doc, "DAFTAR ISI", page_break=True)
    toc_line(doc, "KATA PENGANTAR", PR['kp'], bab=True)
    for b in BAB:
        toc_line(doc, "BAB %s  %s" % (b['no'], b['judul']), PR['bab' + b['no']], bab=True)
        k = 0
        for blk in b['isi']:
            if blk[0] == 'h2':
                k += 1
                toc_line(doc, blk[1], PR.get("s%s_%d" % (b['no'], k), ""), bab=False)
    toc_line(doc, "DAFTAR PUSTAKA", PR['dp'], bab=True)

    for b in BAB:
        heading(doc, "BAB %s" % b['no'], size=13, page_break=True, after=4, color=GOLD_DARK)
        heading(doc, b['judul'], size=16, after=18)
        for blk in b['isi']:
            if blk[0] == 'h2':
                heading(doc, blk[1], size=13, align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=6)
            elif blk[0] == 'akar':
                add_box(doc, blk[1], 'F7EBED', MAROON_HEX, 'B9485A', MAROON)
            elif blk[0] == 'solusi':
                add_box(doc, blk[1], 'FBF4E3', 'A87F23', GOLD_HEX, GOLD_DARK)
            else:
                para(doc, blk[1])

    heading(doc, "DAFTAR PUSTAKA", size=16, page_break=True, after=18)
    for pre, title, post in DAFTAR_PUSTAKA:
        p = doc.add_paragraph(); pf = p.paragraph_format; ls_set(pf)
        pf.space_after = Pt(8); pf.left_indent = Cm(1.0); pf.first_line_indent = Cm(-1.0); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        widow(p)
        r1 = p.add_run(pre); set_font(r1, 12)
        r2 = p.add_run(title); set_font(r2, 12, italic=True)
        r3 = p.add_run(post); set_font(r3, 12)

    full_cover(doc, back, first=False)
    enable_update_fields(doc)
    os.makedirs(os.path.dirname(OUT), exist_ok=True); doc.save(OUT)
    print("SAVED:", OUT, os.path.getsize(OUT), "bytes")

if __name__ == "__main__":
    build()
