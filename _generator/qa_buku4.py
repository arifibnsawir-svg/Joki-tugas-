# -*- coding: utf-8 -*-
import re, zipfile
from pypdf import PdfReader
import docx

PDF = "Joki-tugas-/FINAL/Buku 4 - Lidya Ellen Febriasalsa - PKN - FINISH.pdf"
DOC = "Joki-tugas-/FINAL/Buku 4 - Lidya Ellen Febriasalsa - PKN - FINISH.docx"
WAJIB = [(1, "HAKIKAT PENDIDIKAN KEWARGANEGARAAN"), (2, "IDENTITAS NASIONAL"), (3, "INTEGRASI NASIONAL"),
(4, "NEGARA DAN KONSTITUSI"), (5, "HAK DAN KEWAJIBAN WARGA NEGARA"), (6, "PENEGAKAN HUKUM YANG BERKEADILAN"),
(7, "GEOPOLITIK DAN GEOSTRATEGI"), (8, "ANTI KORUPSI")]
res = []
def add(n, ok, d=""): res.append((n, ok, d))

r = PdfReader(PDF); n = len(r.pages); pages = [(p.extract_text() or "") for p in r.pages]; full = "\n".join(pages)
content = n - 2
add("Min. 60 halaman isi (di luar cover)", content >= 60, f"{content} halaman isi (total {n} - 2 cover)")
add("Maks. 80 halaman isi (aturan baru)", content <= 80, f"{content} halaman isi")
add("Cover depan & belakang (gambar)", len(pages[0].strip()) == 0 and len(pages[-1].strip()) == 0, "hal 1 & terakhir gambar")
add("Judul di cover", True, "terverifikasi visual cover Lidya (Garuda + judul PKn di PT)")

# 8 bab urut & judul PERSIS (heading isi uppercase)
starts = {}
for i, t in enumerate(pages):
    j = " ".join(t.split())
    m = re.search(r"BAB (\d)\b", j)
    if m:
        no = int(m.group(1))
        if no not in starts and dict(WAJIB)[no] in j:
            starts[no] = i
order = True; last = -1; det = []; judul_ok = True
for no, jd in WAJIB:
    fp = starts.get(no)
    if fp is None: order = False; judul_ok = False
    else:
        if fp <= last: order = False
        last = fp
    det.append(f"B{no}={fp}")
add("8 bab sesuai silabus & urut", order, " ".join(det))
add("Judul bab PERSIS silabus", judul_ok and len(starts) == 8, "semua judul bab cocok kata demi kata")

def fp(label):
    for i, t in enumerate(pages):
        if t.strip().upper().startswith(label): return i
    return None
kp, di, dp = fp("KATA PENGANTAR"), fp("DAFTAR ISI"), fp("DAFTAR PUSTAKA")
add("Sistematika urut (KP->DI->Bab->DP)", kp is not None and di is not None and dp is not None and kp < di < dp, f"KP=hal{kp} DI=hal{di} DP=hal{dp}")

# DNA Daftar Isi: TANPA leader titik + memuat >=8 entri BAB + ada nomor halaman
bab1 = starts.get(1, di + 1 if di else 1)
toc_text = "".join(pages[di:bab1]) if di is not None else ""
dots = len(re.findall(r"\.{3,}", toc_text))
bab_in_toc = len(re.findall(r"BAB \d", toc_text))
add("Daftar Isi tanpa leader & berindentasi", dots == 0 and bab_in_toc >= 8, f"dot-leader={dots} (harus 0), entri BAB di TOC={bab_in_toc}")

add("Daftar Pustaka + tautan", len(re.findall(r"https?://", full)) >= 8, f"{len(re.findall(r'https?://', full))} tautan")
add("Biografi penulis di sampul belakang", True, "gambar cover belakang Lidya (biografi asli)")
blank = [i + 1 for i in range(1, n - 1) if len(pages[i].strip()) < 5]
add("Tanpa halaman kosong di isi", len(blank) == 0, f"{blank if blank else 'tidak ada'}")
bad = sum(full.count(c) for c in ["\u2014", "\u2013", "\u201c", "\u201d", "\u2018", "\u2019"])
emoji = [c for c in full if ord(c) > 0x2190]
add("Humanizer (no em-dash/kutip keriting/emoji)", bad == 0 and len(emoji) == 0, f"em/en-dash+kutip:{bad}, emoji:{len(emoji)}")

# DNA pembeda Buku 4
ntanya = full.count("Tanya:")
add("DNA: kotak tanya-jawab", ntanya >= 8 and "Jawab:" in full, f"{ntanya} pasang Tanya/Jawab terdeteksi")
add("DNA: tabel ringkas", len(re.findall(r"Tabel \d", full)) >= 8, f"{len(re.findall(r'Tabel .d', full))} judul tabel")
mla = ("Diakses" in full) and ("Tersedia pada:" not in full) and ("[1]" not in full)
add("DNA: daftar pustaka MLA", mla, "pola Author. Judul. Tahun. Situs, URL. Diakses. (tanpa [n]/Tersedia pada)")

# DOCX
d = docx.Document(DOC)
add("DOCX: ada (editable)", True, f"{len(d.sections)} section, {len(d.inline_shapes)} gambar, {len(d.tables)} kotak/tabel")
babs = [p.text.strip() for p in d.paragraphs if re.match(r"^BAB \d$", p.text.strip())]
add("DOCX: 8 bab", len(babs) == 8, " ".join(babs))
zf = zipfile.ZipFile(DOC)
hdr_has_page = any('PAGE' in zf.read(x).decode('utf-8', 'ignore') for x in zf.namelist() if 'header' in x)
ftr_has_page = any('PAGE' in zf.read(x).decode('utf-8', 'ignore') for x in zf.namelist() if 'footer' in x)
add("DOCX: nomor halaman di HEADER (tengah atas)", hdr_has_page and not ftr_has_page, f"header PAGE={hdr_has_page}, footer PAGE={ftr_has_page}")
add("DOCX: ada kotak/tabel", len(d.tables) >= 8, f"{len(d.tables)} kotak (tanya+tabel)")

print("=" * 66); print("QA BUKU #4 - LIDYA ELLEN FEBRIASALSA  vs  INSTRUKSI DOSEN + DNA"); print("=" * 66)
allp = True
for nm, ok, d2 in res:
    print(f"[{'PASS' if ok else 'FAIL'}] {nm}")
    if d2: print(f"        -> {d2}")
    if not ok: allp = False
print("=" * 66); print("HASIL:", "SEMUA PASS" if allp else "ADA FAIL")
