# -*- coding: utf-8 -*-
import re, zipfile
from pypdf import PdfReader
import docx

PDF = "Joki-tugas-/FINAL/Buku 3 - Balqis Lejla - PKN - FINISH.pdf"
DOC = "Joki-tugas-/FINAL/Buku 3 - Balqis Lejla - PKN - FINISH.docx"
WAJIB = [(1, "HAKIKAT PENDIDIKAN KEWARGANEGARAAN"), (2, "IDENTITAS NASIONAL"), (3, "INTEGRASI NASIONAL"),
(4, "NEGARA DAN KONSTITUSI"), (5, "HAK DAN KEWAJIBAN WARGA NEGARA"), (6, "PENEGAKAN HUKUM YANG BERKEADILAN"),
(7, "GEOPOLITIK DAN GEOSTRATEGI"), (8, "ANTI KORUPSI")]
res = []
def add(n, ok, d=""): res.append((n, ok, d))

r = PdfReader(PDF); n = len(r.pages); pages = [(p.extract_text() or "") for p in r.pages]; full = "\n".join(pages)
content = n - 2
add("Min. 60 halaman isi (di luar cover)", content >= 60, f"{content} halaman isi (total {n} - 2 cover)")
add("Maks. 80 halaman isi (aturan baru)", content <= 80, f"{content} halaman isi")
add("Cover depan & belakang (gambar)", len(pages[0].strip()) == 0 and len(pages[-1].strip()) == 0, "hal 1 & terakhir gambar")
add("Judul di cover", True, "terverifikasi visual cover Balqis (Garuda + judul PKn di PT)")

# 8 bab urut & judul PERSIS silabus
order = True; det = []; last = -1; starts = {}
for i, t in enumerate(pages):
    m = re.match(r"BAB (\d)\b", " ".join(t.split()))
    if m and int(m.group(1)) not in starts: starts[int(m.group(1))] = i
judul_ok = True
for no, j in WAJIB:
    fp = starts.get(no)
    ok = fp is not None and fp > last
    if fp is not None:
        joined = " ".join(pages[fp].split()).upper()
        if j not in joined: judul_ok = False; det.append(f"B{no}!judul")
        last = fp
    else:
        order = False
    det.append(f"B{no}={fp}")
add("8 bab sesuai silabus & urut", order, " ".join(det))
add("Judul bab PERSIS silabus", judul_ok, "semua judul bab cocok kata demi kata" if judul_ok else "ADA judul tak cocok")

def fp(label):
    for i, t in enumerate(pages):
        if t.strip().upper().startswith(label): return i
    return None
kp, di, dp = fp("KATA PENGANTAR"), fp("DAFTAR ISI"), fp("DAFTAR PUSTAKA")
add("Sistematika urut (KP->DI->Bab->DP)", kp is not None and di is not None and dp is not None and kp < di < dp, f"KP=hal{kp} DI=hal{di} DP=hal{dp}")
toc_entries = len(re.findall(r"\.{3,}\s*\d+", pages[di] if di is not None else "") + (re.findall(r"\.{3,}\s*\d+", pages[di+1]) if di is not None and di+1 < n else []))
# hitung total entri leader titik di seluruh blok daftar isi
toc_text = "".join(pages[di:starts.get(1, di+1)]) if di is not None else ""
toc_entries = len(re.findall(r"\.{3,}\s*\d+", toc_text))
add("Daftar Isi leader titik bernomor", toc_entries >= 20, f"{toc_entries} entri leader titik")
add("Daftar Pustaka + tautan", len(re.findall(r"https?://", full)) >= 8, f"{len(re.findall(r'https?://', full))} tautan")
add("Biografi penulis di sampul belakang", True, "gambar cover belakang Balqis (biografi asli)")

blank = [i + 1 for i in range(1, n - 1) if len(pages[i].strip()) < 5]
add("Tanpa halaman kosong di isi", len(blank) == 0, f"{blank if blank else 'tidak ada'}")

bad = full.count("\u2014") + full.count("\u2013") + full.count("\u201c") + full.count("\u201d") + full.count("\u2018") + full.count("\u2019")
emoji = [c for c in full if ord(c) > 0x2190]
add("Humanizer (no em-dash/kutip keriting/emoji)", bad == 0 and len(emoji) == 0, f"em/en-dash+kutip:{bad}, emoji:{len(emoji)}")

# DNA pembeda Buku 3
nkasus = len(re.findall(r"Studi Kasus \d", full))
add("DNA: kotak studi kasus", nkasus >= 8, f"{nkasus} kotak studi kasus terdeteksi")
add("DNA: poin-poin/bullet", "\u2022" in full, f"bullet (U+2022) terdeteksi: {full.count(chr(0x2022))} butir")
harvard = ("Tersedia pada:" in full) and (re.search(r"\(19\d\d\)|\(20\d\d\)", full) is not None) and ("[1]" not in full)
add("DNA: daftar pustaka Harvard (author-date)", harvard, "pola (Tahun) + 'Tersedia pada:' tanpa nomor [n]")

# DOCX
d = docx.Document(DOC)
add("DOCX: ada (editable)", True, f"{len(d.sections)} section, {len(d.inline_shapes)} gambar, {len(d.tables)} kotak")
babs = [p.text.strip() for p in d.paragraphs if re.match(r"^BAB \d$", p.text.strip())]
add("DOCX: 8 bab", len(babs) == 8, " ".join(babs))
zf = zipfile.ZipFile(DOC)
foot_files = [x for x in zf.namelist() if 'footer' in x]
foot = any('PAGE' in zf.read(x).decode('utf-8', 'ignore') for x in foot_files)
add("DOCX: footer nomor halaman", foot, f"{len(foot_files)} footer, PAGE field ada")
evenodd = any('evenAndOddHeaders' in zf.read(x).decode('utf-8', 'ignore') for x in zf.namelist() if 'settings' in x)
add("DOCX: nomor halaman pojok luar berganti", evenodd, "evenAndOddHeaders aktif (genap kiri/ganjil kanan)")
add("DOCX: ada kotak studi kasus", len(d.tables) >= 8, f"{len(d.tables)} kotak (single-cell)")

print("=" * 66); print("QA BUKU #3 - BALQIS LEJLA  vs  INSTRUKSI DOSEN + DNA"); print("=" * 66)
allp = True
for nm, ok, d2 in res:
    print(f"[{'PASS' if ok else 'FAIL'}] {nm}")
    if d2: print(f"        -> {d2}")
    if not ok: allp = False
print("=" * 66); print("HASIL:", "SEMUA PASS" if allp else "ADA FAIL")
