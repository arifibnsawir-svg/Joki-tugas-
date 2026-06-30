# -*- coding: utf-8 -*-
"""DOCX editable Modul Gaya Belajar (Kelompok 10) - mirror PDF premium.
Format dosen: A4, Cambria 12, spasi 1.5, justify, margin kiri4/kanan3/atas3/bawah3, nomor halaman Arab.
Cover dirasterisasi dari PDF (pymupdf) lalu disisipkan full-page. Tabel berwarna, gambar disisipkan,
kotak (single-cell shaded), daftar isi/tabel/gambar manual dari page_real.json.
"""
import os, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import konten as K

OUTDIR = "Joki-tugas-/MODUL_FINAL"
OUT = os.path.join(OUTDIR, "Modul Gaya Belajar - Kelompok 10 - BnP - FINAL.docx")
PDF = os.path.join(OUTDIR, "Modul Gaya Belajar - Kelompok 10 - BnP - FINAL.pdf")
ASSET = "Joki-tugas-/_modul/gaya-belajar/assets"
PR = json.load(open("Joki-tugas-/_modul/gaya-belajar/page_real.json"))
FONT = "Cambria"
TW = 14.0  # lebar teks cm (21 - 4 - 3)
LS = 1.5
INDIGO = RGBColor(0x1F, 0x3C, 0x88)
GOLDD = RGBColor(0x7A, 0x5B, 0x12)
GREY = RGBColor(0x5A, 0x64, 0x77)
BOX = {  # kind -> (garis/judul hex, latar hex, judul RGB)
    "tujuan": ("1F3C88", "EEF1FB", RGBColor(0x1F, 0x3C, 0x88)),
    "contoh": ("0E7490", "E6F6FA", RGBColor(0x0E, 0x74, 0x90)),
    "aktivitas": ("EA7317", "FDF1E6", RGBColor(0xC2, 0x5E, 0x10)),
    "ringkasan": ("15A34A", "E9F7EE", RGBColor(0x10, 0x7A, 0x37)),
    "tips": ("7C3AED", "F3EDFD", RGBColor(0x60, 0x2A, 0xC2)),
}

def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT; run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    if color is not None: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    rf.set(qn('w:ascii'), FONT); rf.set(qn('w:hAnsi'), FONT); rf.set(qn('w:eastAsia'), FONT)

def ls_set(pf, val=LS):
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = val

def flag(p, tag):
    e = OxmlElement('w:' + tag); e.set(qn('w:val'), 'true'); p._p.get_or_add_pPr().append(e)
def widow(p): flag(p, 'widowControl')

def shade(el, fill):
    # el: paragraph or cell ; pasang shading
    pr_ = el._p.get_or_add_pPr() if hasattr(el, "_p") else el._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill); pr_.append(sh)

def cell_shade(cell, fill):
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(sh)

def page_field(p):
    for spec in ['begin', None, 'separate', '1', 'end']:
        run = p.add_run()
        if spec == '1':
            run.text = '1'
        elif spec is None:
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '
            run._r.append(it)
        else:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), spec); run._r.append(fc)
        set_font(run, 11, color=INDIGO)

def enable_update(doc):
    s = doc.settings.element
    if s.find(qn('w:updateFields')) is None:
        u = OxmlElement('w:updateFields'); u.set(qn('w:val'), 'true'); s.append(u)

def pgstart(section, start=1):
    sp = section._sectPr; pg = sp.find(qn('w:pgNumType'))
    if pg is None: pg = OxmlElement('w:pgNumType'); sp.append(pg)
    pg.set(qn('w:start'), str(start))

def para(doc, text, size=12, after=8, before=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, color=None, indent=None):
    p = doc.add_paragraph(); pf = p.paragraph_format; p.alignment = align; ls_set(pf)
    pf.space_after = Pt(after); pf.space_before = Pt(before); widow(p)
    if indent is not None:
        pf.left_indent = Cm(indent); pf.first_line_indent = Cm(-indent)
    r = p.add_run(text); set_font(r, size, bold=bold, color=color); return p

def h1(doc, text):
    p = doc.add_paragraph(); pf = p.paragraph_format; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ls_set(pf); pf.space_after = Pt(14); pf.page_break_before = True; pf.keep_with_next = True; widow(p)
    bd = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '18'); bt.set(qn('w:space'), '6'); bt.set(qn('w:color'), 'C8A24A')
    bd.append(bt); p._p.get_or_add_pPr().append(bd)
    r = p.add_run(text); set_font(r, 16, bold=True, color=INDIGO); return p

def h2(doc, text):
    p = doc.add_paragraph(); pf = p.paragraph_format; ls_set(pf)
    pf.space_before = Pt(11); pf.space_after = Pt(5); pf.keep_with_next = True; widow(p)
    bd = OxmlElement('w:pBdr'); lf = OxmlElement('w:left')
    lf.set(qn('w:val'), 'single'); lf.set(qn('w:sz'), '24'); lf.set(qn('w:space'), '8'); lf.set(qn('w:color'), 'C8A24A')
    bd.append(lf); p._p.get_or_add_pPr().append(bd)
    pf.left_indent = Cm(0.25)
    r = p.add_run(text); set_font(r, 13, bold=True, color=INDIGO); return p

def nlist(doc, items, ordered=True):
    for i, it in enumerate(items, 1):
        p = doc.add_paragraph(); pf = p.paragraph_format; ls_set(pf); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.left_indent = Cm(1.0); pf.first_line_indent = Cm(-0.5); pf.space_after = Pt(4); widow(p)
        prefix = ("%d. " % i) if ordered else "\u2022 "
        r = p.add_run(prefix + it); set_font(r, 12)

def cantsplit(tbl):
    for row in tbl.rows:
        tr = row._tr.get_or_add_trPr(); cs = OxmlElement('w:cantSplit'); cs.set(qn('w:val'), 'true'); tr.append(cs)

def box(doc, d):
    line, bg, tcol = BOX[d["kind"]]
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # border: kiri tebal warna, sisanya tipis
    tblPr = t._tbl.tblPr; brd = OxmlElement('w:tblBorders')
    for edge, sz, col in [('top', '6', 'DFE3EE'), ('bottom', '6', 'DFE3EE'), ('right', '6', 'DFE3EE'),
                          ('left', '30', line), ('insideH', '6', 'DFE3EE'), ('insideV', '6', 'DFE3EE')]:
        e = OxmlElement('w:' + edge); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz); e.set(qn('w:color'), col); brd.append(e)
    tblPr.append(brd)
    cell = t.rows[0].cells[0]; cell_shade(cell, bg)
    p0 = cell.paragraphs[0]; p0.paragraph_format.space_after = Pt(5); ls_set(p0.paragraph_format, 1.3); widow(p0)
    r = p0.add_run(d["judul"]); set_font(r, 12, bold=True, color=tcol)
    for blk in d["body"]:
        if blk[0] == "p":
            pp = cell.add_paragraph(); pf = pp.paragraph_format; ls_set(pf, 1.4); pf.space_after = Pt(5)
            pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; widow(pp)
            rr = pp.add_run(blk[1]); set_font(rr, 11.5)
        elif blk[0] in ("nlist", "blist"):
            for i, it in enumerate(blk[1], 1):
                pp = cell.add_paragraph(); pf = pp.paragraph_format; ls_set(pf, 1.4)
                pf.left_indent = Cm(0.9); pf.first_line_indent = Cm(-0.5); pf.space_after = Pt(3)
                pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; widow(pp)
                pre = ("%d. " % i) if blk[0] == "nlist" else "\u2022 "
                rr = pp.add_run(pre + it); set_font(rr, 11.5)
    cantsplit(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def tbl_caption(doc, text):
    p = doc.add_paragraph(); pf = p.paragraph_format; ls_set(pf, 1.2)
    pf.space_before = Pt(8); pf.space_after = Pt(3); pf.keep_with_next = True; widow(p)
    r = p.add_run(text); set_font(r, 11, bold=True, color=INDIGO)

def data_table(doc, head, rows, sumber=None):
    t = doc.add_table(rows=1, cols=len(head)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hc = t.rows[0].cells
    for j, htxt in enumerate(head):
        cell_shade(hc[j], "1F3C88")
        p = hc[j].paragraphs[0]; ls_set(p.paragraph_format, 1.2); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(htxt); set_font(r, 10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            if ri % 2 == 1: cell_shade(cells[j], "F1F4FB")
            p = cells[j].paragraphs[0]; ls_set(p.paragraph_format, 1.25); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val if val else ""); set_font(r, 10.5)
    cantsplit(t)
    if sumber:
        p = doc.add_paragraph(); pf = p.paragraph_format; ls_set(pf, 1.2); pf.space_after = Pt(9); widow(p)
        r = p.add_run("Sumber: " + sumber); set_font(r, 9.5, italic=True, color=GREY)

def figure(doc, d):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(os.path.join(ASSET, d["file"]), width=Cm(14.5))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; ls_set(cap.paragraph_format, 1.2)
    cap.paragraph_format.space_after = Pt(2); cap.paragraph_format.keep_with_next = True; widow(cap)
    r = cap.add_run("Gambar %s  %s" % (d["no"], d["judul"])); set_font(r, 11, bold=True, color=INDIGO)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER; ls_set(s.paragraph_format, 1.2)
    s.paragraph_format.space_after = Pt(9); widow(s)
    rr = s.add_run("Sumber: " + d.get("sumber", "")); set_font(rr, 9.5, italic=True, color=GREY)

def render_blocks(doc, blocks):
    for blk in blocks:
        k = blk[0]
        if k == "p": para(doc, blk[1])
        elif k == "h2": h2(doc, blk[1])
        elif k == "nlist": nlist(doc, blk[1], True)
        elif k == "blist": nlist(doc, blk[1], False)
        elif k == "box": box(doc, blk[1])
        elif k == "table":
            d = blk[1]; tbl_caption(doc, "Tabel %s  %s" % (d["no"], d["judul"]))
            data_table(doc, d["head"], d["rows"], d.get("sumber"))
        elif k == "fig": figure(doc, blk[1])

def toc_row(doc, title, page, bab=False, sub=False):
    p = doc.add_paragraph(); pf = p.paragraph_format; ls_set(pf, 1.25)
    pf.space_after = Pt(2); pf.space_before = Pt(6 if bab else 1); widow(p)
    if sub: pf.left_indent = Cm(0.9)
    pf.tab_stops.add_tab_stop(Cm(TW), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(title + "\t" + str(page))
    set_font(r, 11.5, bold=bab, color=INDIGO if bab else None)

def cover_image(doc):
    import fitz
    d = fitz.open(PDF)
    out = os.path.join(ASSET, "cover_render.png")
    d[0].get_pixmap(matrix=fitz.Matrix(200/72, 200/72)).save(out)
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Cm(0)
    sec.header_distance = Cm(0); sec.footer_distance = Cm(0)
    sec.header.is_linked_to_previous = False; sec.footer.is_linked_to_previous = False
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
    p.add_run().add_picture(out, width=Cm(21), height=Cm(29.7))

def build():
    doc = Document()
    doc.styles['Normal'].font.name = FONT; doc.styles['Normal'].font.size = Pt(12)
    cover_image(doc)
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width = Cm(21); body.page_height = Cm(29.7)
    body.top_margin = Cm(3); body.bottom_margin = Cm(3); body.left_margin = Cm(4); body.right_margin = Cm(3)
    pgstart(body, 1)
    body.header.is_linked_to_previous = False
    body.footer.is_linked_to_previous = False
    fp = body.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER; page_field(fp)

    # KATA PENGANTAR
    h1(doc, "KATA PENGANTAR")
    for t in K.KATA_PENGANTAR: para(doc, t)
    para(doc, "Jakarta, Juni 2026\nTim Penyusun,\n" + K.KELOMPOK, align=WD_ALIGN_PARAGRAPH.RIGHT, before=8)

    # DAFTAR ISI
    h1(doc, "DAFTAR ISI")
    toc_row(doc, "Kata Pengantar", PR["kp"])
    toc_row(doc, "Daftar Tabel", PR["dtab"])
    toc_row(doc, "Daftar Gambar", PR["dgam"])
    toc_row(doc, "Pendahuluan", PR["pend"])
    for bab in K.BAB:
        toc_row(doc, "BAB %s  %s" % (bab["no"], bab["judul"]), PR["bab_%s" % bab["no"]], bab=True)
        k = 0
        for blk in bab["isi"]:
            if blk[0] == "h2":
                k += 1
                toc_row(doc, blk[1], PR.get("s_%s_%d" % (bab["no"], k), ""), sub=True)
    toc_row(doc, "Refleksi", PR["refleksi"], bab=True)
    toc_row(doc, "Evaluasi", PR["evaluasi"], bab=True)
    toc_row(doc, "Daftar Pustaka", PR["pustaka"], bab=True)
    toc_row(doc, "Penutup", PR["penutup"], bab=True)

    # DAFTAR TABEL
    h1(doc, "DAFTAR TABEL")
    for bab in K.BAB:
        for blk in bab["isi"]:
            if blk[0] == "table":
                d = blk[1]; toc_row(doc, "Tabel %s  %s" % (d["no"], d["judul"]), PR.get("tb_" + d["no"].replace(".", "_"), ""))
    # DAFTAR GAMBAR
    h1(doc, "DAFTAR GAMBAR")
    for bab in K.BAB:
        for blk in bab["isi"]:
            if blk[0] == "fig":
                d = blk[1]; toc_row(doc, "Gambar %s  %s" % (d["no"], d["judul"]), PR.get("fg_" + d["no"].replace(".", "_"), ""))

    # PENDAHULUAN
    h1(doc, "PENDAHULUAN")
    render_blocks(doc, K.PENDAHULUAN)

    # BAB I-III
    for bab in K.BAB:
        # babhead (kotak indigo)
        t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.rows[0].paragraphs if False else None
        cell = t.rows[0].cells[0]; cell_shade(cell, "1F3C88")
        cell.paragraphs[0].paragraph_format.page_break_before = True
        p1 = cell.paragraphs[0]; ls_set(p1.paragraph_format, 1.1); p1.paragraph_format.space_after = Pt(2)
        r = p1.add_run("BAB %s" % bab["no"]); set_font(r, 12, bold=True, color=RGBColor(0xC8, 0xA2, 0x4A))
        p2 = cell.add_paragraph(); ls_set(p2.paragraph_format, 1.15); p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(bab["judul"]); set_font(r2, 16, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        cantsplit(t)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        render_blocks(doc, bab["isi"])

    # REFLEKSI
    R = K.REFLEKSI
    h1(doc, "REFLEKSI")
    para(doc, R["intro"]); nlist(doc, R["pertanyaan"], True)
    tbl_caption(doc, R["tabel_judul"]); para(doc, R["tabel_intro"])
    data_table(doc, R["tabel"]["head"], R["tabel"]["rows"])

    # EVALUASI
    E = K.EVALUASI
    h1(doc, "EVALUASI")
    para(doc, E["intro"])
    h2(doc, "A. Pilihan Ganda"); para(doc, "Pilihlah satu jawaban yang paling tepat.")
    letters = "ABCD"
    for idx, (q, opts, _) in enumerate(E["pg"], 1):
        para(doc, "%d. %s" % (idx, q), after=2, indent=0.6)
        for j, o in enumerate(opts):
            p = doc.add_paragraph(); pf = p.paragraph_format; ls_set(pf, 1.3); pf.left_indent = Cm(1.3); pf.space_after = Pt(1)
            r = p.add_run("%s. %s" % (letters[j], o)); set_font(r, 12)
    h2(doc, "B. Benar atau Salah"); para(doc, "Tentukan apakah pernyataan berikut benar (B) atau salah (S).")
    nlist(doc, ["%s ( ... )" % s for s, _ in E["bs"]], True)
    h2(doc, "C. Esai Singkat"); para(doc, "Jawablah dengan jelas dan ringkas dalam tiga sampai lima kalimat.")
    nlist(doc, E["esai"], True)
    pgk = "  ".join("%d.%s" % (i + 1, letters[k]) for i, (_, _, k) in enumerate(E["pg"]))
    bsk = "  ".join("%d.%s" % (i + 1, "B" if v else "S") for i, (_, v) in enumerate(E["bs"]))
    box(doc, {"kind": "ringkasan", "judul": "Kunci Jawaban", "body": [
        ("p", "A. Pilihan Ganda: " + pgk),
        ("p", "B. Benar/Salah: " + bsk),
        ("p", "C. Esai: " + E["pedoman_esai"]),
    ]})

    # DAFTAR PUSTAKA
    h1(doc, "DAFTAR PUSTAKA")
    for pre, judul, post, url in K.DAFTAR_PUSTAKA:
        p = doc.add_paragraph(); pf = p.paragraph_format; ls_set(pf); pf.space_after = Pt(7)
        pf.left_indent = Cm(1.0); pf.first_line_indent = Cm(-1.0); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; widow(p)
        r1 = p.add_run(pre); set_font(r1, 12)
        r2 = p.add_run(judul); set_font(r2, 12, italic=True)
        r3 = p.add_run(post + " "); set_font(r3, 12)
        r4 = p.add_run(url); set_font(r4, 11, color=INDIGO)

    # PENUTUP
    h1(doc, "PENUTUP")
    for t in K.PENUTUP: para(doc, t)
    pm = doc.add_paragraph(); pf = pm.paragraph_format; ls_set(pf, 1.3); pf.space_before = Pt(8); pf.space_after = Pt(8)
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bd = OxmlElement('w:pBdr'); lf = OxmlElement('w:left')
    lf.set(qn('w:val'), 'single'); lf.set(qn('w:sz'), '24'); lf.set(qn('w:space'), '10'); lf.set(qn('w:color'), 'C8A24A')
    bd.append(lf); pm._p.get_or_add_pPr().append(bd); shade(pm, "EEF1FB")
    r = pm.add_run(K.MOTTO); set_font(r, 12.5, italic=True, color=INDIGO)
    para(doc, "Jakarta, Juni 2026\nTim Penyusun, " + K.KELOMPOK, align=WD_ALIGN_PARAGRAPH.RIGHT, before=8)

    enable_update(doc)
    os.makedirs(OUTDIR, exist_ok=True); doc.save(OUT)
    print("SAVED:", OUT, os.path.getsize(OUT), "bytes")

if __name__ == "__main__":
    build()
